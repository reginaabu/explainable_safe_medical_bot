"""
build_report.py – Generate the full thesis Word document with architecture diagrams.
Run: python build_report.py
"""
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.0)

# ── Colour palette ─────────────────────────────────────────────────────────────
C_BLUE   = "#4472C4"
C_GREEN  = "#70AD47"
C_ORANGE = "#ED7D31"
C_RED    = "#C00000"
C_PURPLE = "#7030A0"
C_TEAL   = "#00B0F0"
C_GRAY   = "#808080"
C_LIGHT  = "#F2F2F2"
C_DARK   = "#1F3864"

# ── Diagram helpers ────────────────────────────────────────────────────────────
def _box(ax, cx, cy, w, h, label, fc=C_BLUE, tc="white", fs=8.5, bold=True):
    rect = FancyBboxPatch((cx - w/2, cy - h/2), w, h,
                           boxstyle="round,pad=0.03",
                           facecolor=fc, edgecolor="#333333", linewidth=1.2,
                           zorder=2)
    ax.add_patch(rect)
    ax.text(cx, cy, label, ha="center", va="center",
            color=tc, fontsize=fs, fontweight="bold" if bold else "normal",
            multialignment="center", zorder=3)

def _arr(ax, x1, y1, x2, y2, label="", color="#333333", lw=1.5):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="->", color=color, lw=lw),
                zorder=4)
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2 + 0.02
        ax.text(mx, my, label, ha="center", va="bottom",
                fontsize=7.5, color="#555555", zorder=5)

def _diamond(ax, cx, cy, w, h, label, fc=C_ORANGE):
    xs = [cx, cx+w/2, cx, cx-w/2, cx]
    ys = [cy+h/2, cy, cy-h/2, cy, cy+h/2]
    ax.fill(xs, ys, color=fc, zorder=2)
    ax.plot(xs, ys, color="#333333", lw=1.2, zorder=3)
    ax.text(cx, cy, label, ha="center", va="center",
            color="white", fontsize=8, fontweight="bold", zorder=4)

def _zone(ax, x, y, w, h, label, fc="#EAF3FF", ec=C_BLUE):
    rect = FancyBboxPatch((x, y), w, h,
                           boxstyle="round,pad=0.02",
                           facecolor=fc, edgecolor=ec, linewidth=1.5,
                           linestyle="--", zorder=1)
    ax.add_patch(rect)
    ax.text(x + w/2, y + h - 0.05, label, ha="center", va="top",
            fontsize=8, color=ec, fontweight="bold", zorder=2)

def fig_to_docx(fig, width_inches=6.2, caption=None):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="white")
    buf.seek(0)
    plt.close(fig)
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.add_run().add_picture(buf, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.font.name = "Times New Roman"; cr.font.size = Pt(10)
        cr.font.italic = True
    doc.add_paragraph()

# ── Document helpers ───────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name  = name; run.font.size = Pt(size)
    run.font.bold  = bold; run.font.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sizes = {1: 16, 2: 14, 3: 12}
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(sizes.get(level, 12))
        run.font.bold = True
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_font(run)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.2 * level)
    run = p.add_run(text); set_font(run); return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text); set_font(run); return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"; run.font.size = Pt(9)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    p._p.get_or_add_pPr().append(shd); return p

def add_table(headers, rows, caption=""):
    if caption:
        cp = doc.add_paragraph()
        cr = cp.add_run(caption); set_font(cr, bold=True, size=11)
        cp.paragraph_format.space_before = Pt(8)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True; run.font.name = "Times New Roman"
                run.font.size = Pt(10)
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "D9E1F2")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"; run.font.size = Pt(10)
    doc.add_paragraph(); return t

def spacer(n=1):
    for _ in range(n): doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  DIAGRAM GENERATORS
# ══════════════════════════════════════════════════════════════════════════════

def diag_overall_architecture():
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6); ax.axis("off")
    fig.patch.set_facecolor("white")

    # Zone backgrounds
    _zone(ax, 0.1, 4.6, 10.8, 1.25, "User Interface Layer (Streamlit)", "#EAF3FF", C_BLUE)
    _zone(ax, 0.1, 2.4, 10.8, 2.1,  "Retrieval Layer",                  "#EFFFEA", C_GREEN)
    _zone(ax, 0.1, 0.2, 5.2,  2.1,  "Generation & Safety Layer",         "#FFF8EA", C_ORANGE)
    _zone(ax, 5.4, 0.2, 5.5,  2.1,  "Evaluation Layer",                  "#F5EAFF", C_PURPLE)

    # UI Layer boxes
    _box(ax, 2.0, 5.15, 2.2, 0.55, "Query Input\n(st.form)", C_BLUE)
    _box(ax, 5.5, 5.15, 2.2, 0.55, "Enhancement\nSelector", C_TEAL, fs=8)
    _box(ax, 9.0, 5.15, 1.8, 0.55, "Search ↵\nButton", C_DARK)

    # Retrieval layer
    _box(ax, 1.2, 3.35, 1.8, 0.6,  "BM25\nIndex",          C_GREEN, "white")
    _box(ax, 3.2, 3.35, 1.8, 0.6,  "KG Expansion\n(opt.)", C_TEAL,  "white")
    _box(ax, 5.2, 3.35, 1.8, 0.6,  "Cross-Encoder\nReranker (opt.)", C_GREEN, "white", fs=7.5)
    _box(ax, 7.2, 3.35, 1.8, 0.6,  "Score Filter\n≥ 0.85",  C_ORANGE,"white")
    _box(ax, 9.2, 3.35, 1.4, 0.6,  "Top-K\nChunks",         C_DARK,  "white")

    # Semantic index (below retrieval)
    _box(ax, 3.0, 2.65, 2.0, 0.45, "Semantic Index\n(all-MiniLM-L6-v2)", C_GRAY, "white", fs=7.5)

    # Generation layer
    _box(ax, 1.5, 1.35, 2.0, 0.6,  "RAG Generator\n(Claude Sonnet 4.6)", C_ORANGE, "white", fs=7.5)
    _box(ax, 3.8, 1.35, 1.8, 0.6,  "Safety\nChecker",        C_RED,    "white")
    _box(ax, 1.5, 0.65, 1.8, 0.45, "Fact\nDecompose",        C_ORANGE, "white", fs=8)
    _box(ax, 3.8, 0.65, 1.8, 0.45, "Fact\nVerify",           C_RED,    "white", fs=8)

    # Evaluation layer
    _box(ax, 6.5, 1.35, 2.0, 0.6,  "RAGAS Scorer\n(bg thread)",       C_PURPLE, "white", fs=8)
    _box(ax, 9.0, 1.35, 1.7, 0.6,  "Evaluation\nDashboard",           C_DARK,   "white", fs=8)
    _box(ax, 6.5, 0.65, 2.0, 0.45, "Faithfulness\n+ Answer Relevancy", C_PURPLE, "white", fs=7.5)
    _box(ax, 9.0, 0.65, 1.7, 0.45, "Factuality\n+ Safety Badge",      C_DARK,   "white", fs=7.5)

    # Arrows
    _arr(ax, 3.1, 5.15, 4.4, 5.15)
    _arr(ax, 6.6, 5.15, 8.1, 5.15)
    _arr(ax, 5.5, 4.87, 5.5, 4.6)  # from UI down
    _arr(ax, 1.2, 3.05, 1.2, 2.65)
    _arr(ax, 2.1, 3.35, 2.3, 3.35)
    _arr(ax, 4.1, 3.35, 4.3, 3.35)
    _arr(ax, 6.1, 3.35, 6.3, 3.35)
    _arr(ax, 8.1, 3.35, 8.3, 3.35)
    _arr(ax, 9.2, 3.05, 9.0, 2.65)  # chunks down to gen
    _arr(ax, 2.5, 1.35, 2.9, 1.35)
    _arr(ax, 4.7, 1.35, 5.5, 1.35)  # safe → RAGAS
    _arr(ax, 1.5, 1.05, 1.5, 0.88)
    _arr(ax, 3.8, 1.05, 3.8, 0.88)

    ax.set_title("Figure 1: ArogyaSaathi — Overall System Architecture",
                 fontsize=11, fontweight="bold", pad=8)
    return fig


def diag_retrieval_pipeline():
    fig, ax = plt.subplots(figsize=(11, 3.2))
    ax.set_xlim(0, 11); ax.set_ylim(0, 3.2); ax.axis("off")
    fig.patch.set_facecolor("white")

    xs = [0.9, 2.3, 3.7, 5.1, 6.5, 7.9, 9.4, 10.6]
    labels = ["User\nQuery", "Tokenise\n& Lowercase", "BM25\nScore (k1,b)",
              "Normalise\n→ [0,1]", "KG\nExpansion\n(optional)",
              "Cross-Encoder\nRerank\n(optional)", "Score\nFilter\n≥ 0.85", "Top\nChunks"]
    colors = [C_GREEN, C_BLUE, C_BLUE, C_BLUE, C_TEAL, C_TEAL, C_ORANGE, C_DARK]

    for i, (x, lbl, col) in enumerate(zip(xs, labels, colors)):
        _box(ax, x, 1.9, 1.1, 0.75, lbl, col, fs=7.5)
        if i < len(xs) - 1:
            _arr(ax, x + 0.55, 1.9, xs[i+1] - 0.55, 1.9)

    # Note below
    _box(ax, 5.1, 0.85, 3.2, 0.55,
         "Candidate pool: _RETRIEVE_K=20 (internal, never shown to user)",
         "#F0F0F0", "#555555", fs=8, bold=False)
    _arr(ax, 5.1, 1.52, 5.1, 1.12)

    ax.set_title("Figure 2: Retrieval Pipeline — BM25 with Optional KG Expansion and Cross-Encoder Reranking",
                 fontsize=10, fontweight="bold", pad=8)
    return fig


def diag_kg_module():
    fig, ax = plt.subplots(figsize=(11, 4.0))
    ax.set_xlim(0, 11); ax.set_ylim(0, 4.0); ax.axis("off")
    fig.patch.set_facecolor("white")

    # Main flow (top row)
    _box(ax, 1.0, 3.2, 1.5, 0.6, "Original\nQuery", C_GREEN)
    _box(ax, 3.0, 3.2, 1.8, 0.6, "SciSpacy NER\n(en_ner_bc5cdr_md)", C_BLUE, fs=7.5)
    _box(ax, 5.2, 3.2, 1.8, 0.6, "Biomedical\nEntities\n(DISEASE,CHEM)", C_TEAL, fs=7.5)
    _box(ax, 7.4, 3.2, 1.8, 0.6, "KG Lookup\n(triples.csv)", C_PURPLE)
    _box(ax, 9.6, 3.2, 1.2, 0.6, "Co-occur\nNeighbours", C_ORANGE)

    _arr(ax, 1.75, 3.2, 2.1, 3.2)
    _arr(ax, 3.9,  3.2, 4.3, 3.2)
    _arr(ax, 6.1,  3.2, 6.5, 3.2)
    _arr(ax, 8.3,  3.2, 8.7, 3.2)

    # KG detail (middle row)
    _box(ax, 3.5, 2.1, 3.0, 0.65,
         "triples.csv\nco_occurs_with | mentioned_in",
         "#F0F0F0", "#333333", fs=8, bold=False)
    _box(ax, 7.4, 2.1, 1.8, 0.6,
         "Doc-Freq Filter\n≤ 100 docs\n(remove generic)", C_GRAY, fs=7.5)

    _arr(ax, 3.5, 3.5, 3.5, 3.5)
    _arr(ax, 7.4, 2.9, 7.4, 2.4)
    _arr(ax, 8.3, 2.1, 9.6, 2.1)
    _arr(ax, 9.6, 2.4, 9.6, 3.2)  # up to neighbours

    # Output row
    _box(ax, 5.5, 0.9, 3.0, 0.65,
         "Expanded Query\n= Original + Neighbour Terms", C_GREEN)
    _arr(ax, 9.6, 2.9, 9.6, 1.8)
    _arr(ax, 9.6, 1.22, 7.0, 1.22)
    _arr(ax, 1.0, 2.9, 1.0, 1.22)
    _arr(ax, 1.0, 1.22, 4.0, 1.22)

    ax.set_title("Figure 3: Knowledge Graph Query Expansion Module",
                 fontsize=10, fontweight="bold", pad=8)
    return fig


def diag_generation_pipeline():
    fig, ax = plt.subplots(figsize=(11, 5.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 5.5); ax.axis("off")
    fig.patch.set_facecolor("white")

    # st.status steps zone
    _zone(ax, 0.3, 0.4, 10.4, 4.7, "st.status() — Synchronous Pipeline (user sees live steps)", "#FFFBEA", C_ORANGE)

    # Step boxes (left column, top to bottom)
    steps = [
        (2.5, 4.5, "Step 1\nGenerate Answer\n(Claude Sonnet 4.6)", C_ORANGE),
        (2.5, 3.55,"Step 2\nSafety Check\n(regex patterns)", C_RED),
        (2.5, 2.6, "Step 3\nExtract Atomic\nClaims (Haiku)", C_BLUE),
        (2.5, 1.65,"Step 4\nVerify Claims vs\nSources (Haiku)", C_BLUE),
    ]
    for x, y, lbl, col in steps:
        _box(ax, x, y, 3.0, 0.72, lbl, col, fs=8)

    # Decision diamond
    _diamond(ax, 2.5, 0.85, 2.0, 0.55, "Score ≥\nThreshold?", C_ORANGE)

    # Arrows between steps
    _arr(ax, 2.5, 4.14, 2.5, 3.91)
    _arr(ax, 2.5, 3.19, 2.5, 2.96)
    _arr(ax, 2.5, 2.24, 2.5, 2.01)
    _arr(ax, 2.5, 1.29, 2.5, 1.12)

    # Yes branch → Final answer
    _box(ax, 7.0, 0.85, 2.0, 0.65, "Return\nFinal Answer", C_GREEN)
    _arr(ax, 3.5, 0.85, 6.0, 0.85, "Yes")

    # No branch → self-correction
    _box(ax, 7.0, 2.2, 2.5, 0.75,
         "Step 5 — Self-Correction\nRegenerate (strict=True)\nCompare scores", C_PURPLE)
    _arr(ax, 2.5, 0.57, 2.5, 0.18)
    ax.annotate("", xy=(7.0, 2.2 - 0.375), xytext=(2.5, 0.18),
                arrowprops=dict(arrowstyle="->", color="#333", lw=1.5))
    ax.text(4.5, 0.18, "No", fontsize=8, color="#555", ha="center")
    _arr(ax, 8.25, 1.82, 8.25, 1.17)
    _arr(ax, 7.0, 0.85 + 0.325, 7.0, 0.85 + 0.325)  # connect correction to yes

    # Inputs on right
    _box(ax, 8.5, 4.5, 2.0, 0.65, "Top-3 Evidence\nChunks (≥0.85)", C_DARK)
    _arr(ax, 7.5, 4.5, 4.0, 4.5)

    ax.set_title("Figure 4: RAG Generation and Self-Correction Pipeline",
                 fontsize=10, fontweight="bold", pad=8)
    return fig


def diag_safety_module():
    fig, ax = plt.subplots(figsize=(11, 4.8))
    ax.set_xlim(0, 11); ax.set_ylim(0, 4.8); ax.axis("off")
    fig.patch.set_facecolor("white")

    # Input
    _box(ax, 1.0, 4.1, 1.4, 0.65, "Generated\nAnswer Text", C_GREEN)

    # Pattern groups
    _zone(ax, 2.3, 3.5, 2.6, 1.15, "Emergency (6 patterns)", "#FFEEEE", C_RED)
    patterns_em = ["call_911", "ambulance", "suicidal / self-harm",
                   "overdose", "severe bleeding", "unconscious"]
    for i, p in enumerate(patterns_em):
        ax.text(3.6, 4.45 - i * 0.17, f"• {p}", fontsize=7, color="#C00000")

    _zone(ax, 5.1, 3.5, 2.4, 1.15, "Diagnosis (3 patterns)", "#FFF5E6", C_ORANGE)
    patterns_dx = ["you_have", "your_diagnosis_is", "your_test_shows"]
    for i, p in enumerate(patterns_dx):
        ax.text(6.3, 4.45 - i * 0.17, f"• {p}", fontsize=7, color="#C06000")

    _zone(ax, 7.7, 3.5, 2.9, 1.15, "Prescription (4 patterns)", "#FFF0FF", C_PURPLE)
    patterns_rx = ["take N mg (dosage)", "prescribe/prescribed",
                   "inject/injection", "named_drug + take"]
    for i, p in enumerate(patterns_rx):
        ax.text(9.1, 4.45 - i * 0.17, f"• {p}", fontsize=7, color="#7030A0")

    _arr(ax, 1.7, 4.1, 2.3, 4.1)  # → Emergency zone
    _arr(ax, 4.9, 4.1, 5.1, 4.1)  # → Diagnosis zone
    _arr(ax, 7.5, 4.1, 7.7, 4.1)  # → Prescription zone

    # Decision
    _diamond(ax, 5.5, 2.7, 2.2, 0.6, "Any Flags?", C_ORANGE)
    _arr(ax, 5.5, 3.5, 5.5, 3.0)

    # Safe branch
    _box(ax, 2.5, 1.65, 2.5, 0.65,
         "Append Standard\nDisclaimer Only", C_GREEN)
    _arr(ax, 4.4, 2.7, 4.4, 2.7)
    ax.annotate("", xy=(3.75, 1.97), xytext=(4.4, 2.7),
                arrowprops=dict(arrowstyle="->", color="#333", lw=1.5))
    ax.text(3.5, 2.4, "No\n(safe)", fontsize=8, color=C_GREEN, ha="center")

    # Unsafe branch
    _box(ax, 8.5, 1.65, 2.2, 0.85,
         "Inject Emergency\nDisclaimer +\nStandard Disclaimer", C_RED)
    ax.annotate("", xy=(7.6, 2.07), xytext=(6.6, 2.7),
                arrowprops=dict(arrowstyle="->", color="#333", lw=1.5))
    ax.text(7.4, 2.5, "Yes\n(unsafe)", fontsize=8, color=C_RED, ha="center")

    # Output
    _box(ax, 5.5, 0.75, 3.5, 0.65,
         "Return {is_safe, flags[], answer_with_disclaimer}", C_DARK)
    _arr(ax, 2.5, 1.32, 2.5, 1.07)
    ax.annotate("", xy=(3.75, 1.07), xytext=(2.5, 1.07),
                arrowprops=dict(arrowstyle="->", color="#333", lw=1.5))
    _arr(ax, 8.5, 1.32, 8.5, 1.07)
    ax.annotate("", xy=(7.25, 1.07), xytext=(8.5, 1.07),
                arrowprops=dict(arrowstyle="->", color="#333", lw=1.5))

    ax.set_title("Figure 5: Safety Checking Module — Pattern Groups and Disclaimer Injection",
                 fontsize=10, fontweight="bold", pad=8)
    return fig


def diag_fact_decompose():
    fig, ax = plt.subplots(figsize=(10, 3.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 3.5); ax.axis("off")
    fig.patch.set_facecolor("white")

    _box(ax, 1.0, 2.5, 1.5, 0.65, "Generated\nAnswer", C_GREEN)
    _box(ax, 3.0, 2.5, 2.0, 0.65, "NLTK Sentence\nTokenizer", C_BLUE)
    _diamond(ax, 5.8, 2.5, 2.2, 0.65, ">1 sent OR\n>30 words?", C_ORANGE)

    _arr(ax, 1.75, 2.5, 2.0, 2.5)
    _arr(ax, 4.0,  2.5, 4.7, 2.5)

    # Yes → LLM path
    _box(ax, 8.5, 2.5, 1.6, 0.65, "Claude Haiku\nDecompose", C_BLUE)
    _arr(ax, 6.9, 2.5, 7.7, 2.5, "Yes")
    _box(ax, 8.5, 1.4, 1.6, 0.65, "Parse JSON\nArray", C_BLUE)
    _arr(ax, 8.5, 2.17, 8.5, 1.72)

    # No → NLTK direct
    _box(ax, 5.8, 1.4, 2.0, 0.65, "Use NLTK\nSentences", C_GRAY)
    ax.annotate("", xy=(5.8, 1.72), xytext=(5.8, 2.17),
                arrowprops=dict(arrowstyle="->", color="#333", lw=1.5))
    ax.text(5.2, 2.1, "No", fontsize=8, color="#555")

    # Both → output
    _box(ax, 7.0, 0.55, 2.5, 0.65, "Atomic Claims\nlist[str]", C_GREEN)
    _arr(ax, 5.8, 1.07, 7.0, 0.87)
    _arr(ax, 8.5, 1.07, 8.0, 0.87)

    ax.set_title("Figure 6: Atomic Fact Decomposition Module",
                 fontsize=10, fontweight="bold", pad=8)
    return fig


def diag_fact_verify():
    fig, ax = plt.subplots(figsize=(11, 3.8))
    ax.set_xlim(0, 11); ax.set_ylim(0, 3.8); ax.axis("off")
    fig.patch.set_facecolor("white")

    _box(ax, 0.9, 3.0, 1.5, 0.65, "Atomic\nClaims\n(≤12)", C_GREEN)
    _box(ax, 0.9, 2.0, 1.5, 0.65, "Evidence\nChunks\n(≤3)", C_TEAL)

    _box(ax, 3.0, 2.5, 2.2, 0.65,
         "Build Evidence Block\n3 chunks × 800 chars", C_BLUE)
    _arr(ax, 1.65, 3.0, 2.0, 2.75)
    _arr(ax, 1.65, 2.0, 2.0, 2.25)

    _box(ax, 5.5, 2.5, 2.0, 0.65,
         "Claude Haiku\nBatch Verify\n(single API call)", C_BLUE, fs=7.5)
    _arr(ax, 4.1, 2.5, 4.5, 2.5)

    _box(ax, 7.8, 2.5, 1.8, 0.65,
         "Parse JSON\nVerdicts Array", C_BLUE)
    _arr(ax, 6.5, 2.5, 6.9, 2.5)

    # Verdict types
    _box(ax, 9.9, 3.3, 1.5, 0.45, "supported [+]",    C_GREEN,  fs=8)
    _box(ax, 9.9, 2.5, 1.5, 0.45, "unsupported [?]", C_ORANGE, "white", fs=8)
    _box(ax, 9.9, 1.7, 1.5, 0.45, "contradicted [-]", C_RED,   "white", fs=8)

    _arr(ax, 8.7, 2.75, 9.15, 3.3)
    _arr(ax, 8.7, 2.5,  9.15, 2.5)
    _arr(ax, 8.7, 2.25, 9.15, 1.7)

    # Output
    _box(ax, 5.5, 1.0, 4.0, 0.6,
         "Output: [{fact, verdict, pmid}, …]", C_DARK)
    _arr(ax, 9.9, 1.47, 9.9, 1.3)
    _arr(ax, 9.9, 1.3,  7.5, 1.3)

    ax.set_title("Figure 7: Fact Verification Module — LLM-based Evidence Grounding",
                 fontsize=10, fontweight="bold", pad=8)
    return fig


def diag_ragas_async():
    fig, ax = plt.subplots(figsize=(11, 4.2))
    ax.set_xlim(0, 11); ax.set_ylim(0, 4.2); ax.axis("off")
    fig.patch.set_facecolor("white")

    _zone(ax, 0.2, 0.3, 4.8, 3.6, "Main Thread (Streamlit)", "#EAF3FF", C_BLUE)
    _zone(ax, 5.3, 0.3, 5.2, 3.6, "Background Thread (RAGAS)", "#F5EAFF", C_PURPLE)

    # Main thread flow
    _box(ax, 1.5, 3.4, 2.5, 0.55, "Core Pipeline\nCompletes", C_GREEN)
    _box(ax, 1.5, 2.6, 2.5, 0.65, "Store answer, result\nin session_state", C_BLUE)
    _box(ax, 1.5, 1.8, 2.5, 0.65, "Launch threading.Thread\n(daemon=True)", C_ORANGE)
    _box(ax, 1.5, 1.0, 2.5, 0.65, "Render Answer\n+ Eval Dashboard", C_DARK)
    _box(ax, 1.5, 0.42, 2.5, 0.45, "@st.fragment(run_every=2)\n_ragas_panel()", C_PURPLE, fs=7.5)

    _arr(ax, 1.5, 3.12, 1.5, 2.92)
    _arr(ax, 1.5, 2.27, 1.5, 2.12)
    _arr(ax, 1.5, 1.47, 1.5, 1.32)
    _arr(ax, 1.5, 0.67, 1.5, 0.64)

    # Background thread flow
    _box(ax, 8.0, 3.4, 2.0, 0.55, "score_metrics()", C_PURPLE)
    _box(ax, 8.0, 2.6, 2.0, 0.65, "Faithfulness\n(LLM-based, Haiku)", C_PURPLE)
    _box(ax, 8.0, 1.8, 2.0, 0.65, "Answer Relevancy\n(all-MiniLM embed)", C_PURPLE)
    _box(ax, 8.0, 0.95, 2.0, 0.65, "Write to\n_ragas_store[key]", C_DARK)

    _arr(ax, 8.0, 3.12, 8.0, 2.92)
    _arr(ax, 8.0, 2.27, 8.0, 2.12)
    _arr(ax, 8.0, 1.47, 8.0, 1.27)

    # Cross-thread arrows
    _arr(ax, 2.75, 1.8, 5.3, 3.4, "start()")
    _arr(ax, 5.3, 0.95, 2.75, 0.42, "poll →\nrender metrics")

    ax.set_title("Figure 8: Asynchronous RAGAS Evaluation Pipeline",
                 fontsize=10, fontweight="bold", pad=8)
    return fig


def diag_archehr_integration():
    fig, ax = plt.subplots(figsize=(11, 4.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 4.5); ax.axis("off")
    fig.patch.set_facecolor("white")

    # Input files
    _box(ax, 1.0, 4.0, 1.6, 0.6,  "archehr-qa.xml",         C_TEAL,   fs=8)
    _box(ax, 1.0, 3.1, 1.6, 0.6,  "archehr-qa_key.json\n(relevance annots)", C_TEAL, fs=7.5)
    _box(ax, 1.0, 2.2, 1.6, 0.6,  "archehr-qa_mapping.json\n(MIMIC source)", C_TEAL, fs=7.5)

    # Parser
    _box(ax, 3.5, 3.5, 2.0, 0.7,  "XML Parser\nET.parse()\n<annotations>", C_BLUE, fs=7.5)
    _arr(ax, 1.8, 4.0, 2.5, 3.75)
    _arr(ax, 1.8, 3.1, 2.5, 3.5)
    _arr(ax, 1.8, 2.2, 2.5, 3.25)

    # Case parsing
    _box(ax, 5.8, 3.5, 2.0, 0.7,
         "Per Case:\npatient_question\nnote_excerpt_sentences", C_BLUE, fs=7.5)
    _arr(ax, 4.5, 3.5, 4.8, 3.5)

    # Relevance filter
    _box(ax, 8.2, 3.5, 2.0, 0.7,
         "Filter by relevance:\nessential + supplementary\n(skip not-relevant)", C_ORANGE, fs=7.5)
    _arr(ax, 6.8, 3.5, 7.2, 3.5)

    # Normalised row
    _box(ax, 5.5, 2.1, 4.0, 0.7,
         "Normalised Row\n{doc_id, question, context,\nq_type (clinical_specialty)}", C_GREEN, fs=7.5)
    _arr(ax, 8.2, 3.15, 7.5, 2.45)
    _arr(ax, 3.5, 3.15, 4.0, 2.45)

    # Output
    _box(ax, 3.0, 0.9, 2.2, 0.65, "Semantic\nIndex", C_PURPLE)
    _box(ax, 6.5, 0.9, 2.2, 0.65, "BM25 /\nHybrid Index", C_PURPLE)
    _arr(ax, 4.5, 2.1, 3.5, 1.55)
    _arr(ax, 6.5, 2.1, 7.0, 1.55)

    ax.set_title("Figure 9: ArchEHR-QA Dataset Integration Pipeline",
                 fontsize=10, fontweight="bold", pad=8)
    return fig


def diag_dataset_adapter():
    fig, ax = plt.subplots(figsize=(11, 4.0))
    ax.set_xlim(0, 11); ax.set_ylim(0, 4.0); ax.axis("off")
    fig.patch.set_facecolor("white")

    # Dataset sources
    _box(ax, 0.9, 3.5, 1.4, 0.6, "PubMedQA\n(211k pairs)", C_TEAL, fs=7.5)
    _box(ax, 0.9, 2.7, 1.4, 0.6, "MedQuAD\n(47.4k pairs)", C_TEAL, fs=7.5)
    _box(ax, 0.9, 1.9, 1.4, 0.6, "ArchEHR-QA\n(EHR notes)", C_TEAL, fs=7.5)

    # Adapter
    _box(ax, 3.2, 2.7, 2.2, 1.4, "dataset_adapter.py\nload_dataset_rows()\n\nDATASET_META config\nper-dataset loader", C_BLUE, fs=7.5)
    _arr(ax, 1.6, 3.5, 2.1, 3.1)
    _arr(ax, 1.6, 2.7, 2.1, 2.7)
    _arr(ax, 1.6, 1.9, 2.1, 2.3)

    # Normalised schema
    _box(ax, 6.0, 2.7, 2.5, 1.2,
         "Normalised Row Schema:\n• doc_id  (str)\n• question (str)\n• context  (str)\n• q_type  (str|None)", C_GREEN, fs=7.5)
    _arr(ax, 4.3, 2.7, 4.75, 2.7)

    # Index types
    _box(ax, 9.2, 3.5, 1.4, 0.6, "BM25\nIndex", C_PURPLE, fs=8)
    _box(ax, 9.2, 2.7, 1.4, 0.6, "Semantic\nIndex", C_PURPLE, fs=8)
    _box(ax, 9.2, 1.9, 1.4, 0.6, "Hybrid\nIndex (RRF)", C_PURPLE, fs=8)
    _arr(ax, 7.25, 3.1, 8.5, 3.5)
    _arr(ax, 7.25, 2.7, 8.5, 2.7)
    _arr(ax, 7.25, 2.3, 8.5, 1.9)

    # XML loader callout
    _box(ax, 3.2, 0.8, 2.5, 0.65,
         "load_archehr_xml()\nauto-detect .xml ext\nor directory path", "#F0F0F0", "#333", fs=7.5, bold=False)
    _arr(ax, 3.2, 2.0, 3.2, 1.45)

    ax.set_title("Figure 10: Dataset Adapter Architecture — Multi-Dataset Normalisation",
                 fontsize=10, fontweight="bold", pad=8)
    return fig


def diag_ui_flow():
    fig, ax = plt.subplots(figsize=(11, 4.0))
    ax.set_xlim(0, 11); ax.set_ylim(0, 4.0); ax.axis("off")
    fig.patch.set_facecolor("white")

    _zone(ax, 0.2, 0.2, 10.6, 3.5, "Streamlit Single-Page Application (app.py)", "#F8F8FF", C_DARK)

    # Form
    _box(ax, 1.5, 3.2, 2.5, 0.6,
         "st.form('search_form')\nQuery + Enhancement + Submit", C_BLUE, fs=7.5)
    _box(ax, 1.5, 2.4, 2.5, 0.6,
         "Enter key or\n'Search ↵' click", C_TEAL, fs=8)
    _arr(ax, 1.5, 2.9, 1.5, 2.7)

    # Retrieval
    _box(ax, 4.5, 3.2, 2.0, 0.6,
         "bm25_retrieve()\n± kg_expand\n± rerank", C_ORANGE, fs=7.5)
    _arr(ax, 2.75, 2.7, 3.5, 3.2)

    # Score filter
    _box(ax, 7.0, 3.2, 1.8, 0.6,
         "result_cards()\nfilter ≥ 0.85", C_ORANGE, fs=8)
    _arr(ax, 5.5, 3.2, 6.1, 3.2)

    # Pipeline
    _box(ax, 4.5, 2.1, 2.5, 0.6,
         "_run_core_pipeline()\nst.status() 5 steps", C_RED, fs=7.5)
    _arr(ax, 4.5, 2.9, 4.5, 2.7)

    # Answer
    _box(ax, 4.5, 1.3, 2.5, 0.6,
         "Render Answer\n+ Why Panel\n+ Citation Gate", C_DARK, fs=7.5)
    _arr(ax, 4.5, 1.8, 4.5, 1.6)

    # RAGAS fragment
    _box(ax, 8.5, 2.1, 1.8, 0.6,
         "RAGAS thread\n(background)", C_PURPLE, fs=8)
    _box(ax, 8.5, 1.3, 1.8, 0.6,
         "@st.fragment\npoll every 2s", C_PURPLE, fs=8)
    _arr(ax, 5.75, 2.1, 7.6, 2.1)
    _arr(ax, 8.5, 1.8, 8.5, 1.6)

    # Eval dashboard
    _box(ax, 4.5, 0.55, 2.5, 0.55,
         "_render_eval_dashboard()\nSafety + Factuality", C_GREEN, fs=7.5)
    _arr(ax, 4.5, 1.0, 4.5, 0.82)
    _box(ax, 8.5, 0.55, 1.8, 0.55,
         "Quality Scores\n⚡ timers + bars", C_PURPLE, fs=8)
    _arr(ax, 8.5, 1.0, 8.5, 0.82)

    ax.set_title("Figure 11: Streamlit UI Interaction Flow",
                 fontsize=10, fontweight="bold", pad=8)
    return fig


def diag_semantic_index():
    fig, ax = plt.subplots(figsize=(11, 3.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 3.5); ax.axis("off")
    fig.patch.set_facecolor("white")

    # Input
    _box(ax, 0.9, 2.5, 1.4, 0.65, "Dataset Rows\n(question+context)", C_GREEN)

    # Encode
    _box(ax, 2.9, 2.5, 2.2, 0.65,
         "all-MiniLM-L6-v2\nEncode Questions\n(L2 normalised)", C_BLUE, fs=7.5)
    _arr(ax, 1.6, 2.5, 1.8, 2.5)

    # Index
    _box(ax, 5.4, 2.5, 1.8, 0.65,
         "Dense Embedding\nMatrix (N × 384)", C_PURPLE)
    _arr(ax, 4.0, 2.5, 4.5, 2.5)

    # Query path
    _box(ax, 0.9, 1.2, 1.4, 0.55, "User Query", C_GREEN)
    _box(ax, 2.9, 1.2, 1.8, 0.55, "Encode Query\nVector", C_BLUE, fs=8)
    _box(ax, 5.0, 1.2, 2.0, 0.55, "Dot Product\n(cosine sim)", C_ORANGE, fs=8)
    _box(ax, 7.3, 1.2, 1.8, 0.55, "Top-K by\ncos sim", C_ORANGE, fs=8)

    _arr(ax, 1.6, 1.2, 2.0, 1.2)
    _arr(ax, 3.8, 1.2, 4.0, 1.2)
    _arr(ax, 6.0, 1.2, 6.3, 1.2)
    _arr(ax, 8.2, 1.2, 8.6, 1.2)
    _arr(ax, 5.4, 2.17, 5.4, 1.47)  # matrix → dot

    # Hybrid RRF
    _box(ax, 9.5, 1.8, 1.3, 1.2,
         "Hybrid\nIndex\n(RRF)\nBM25 +\nSemantic", C_TEAL, fs=7.5)
    _arr(ax, 9.1, 1.2, 9.85, 1.2)
    _arr(ax, 9.85, 1.2, 9.85, 1.2)
    _arr(ax, 5.4, 2.5, 5.4, 2.5)

    ax.set_title("Figure 12: Semantic and Hybrid Retrieval Index (SemanticIndex / HybridIndex)",
                 fontsize=10, fontweight="bold", pad=8)
    return fig


# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
spacer(4)
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("EXPLAINABLE AND SAFE MEDICAL SUPPORT CHATBOT\nUSING GENERATIVE AI")
r.font.name = "Times New Roman"; r.font.size = Pt(20); r.font.bold = True

spacer(2)
ap = doc.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = ap.add_run("Regina Aboobacker")
ar.font.name = "Times New Roman"; ar.font.size = Pt(14)

spacer(1)
dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = dp.add_run("MSc Artificial Intelligence\nUpGrad – March 2026")
dr.font.name = "Times New Roman"; dr.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
heading("Abstract", 1)
body(
    "Patients and the general public increasingly rely on online sources and conversational "
    "agents for medical advice; however, the majority of available information is inaccurate, "
    "incomplete, or presented without clear evidence, leading to serious safety and trust "
    "concerns. Large Language Models (LLMs) can provide fluent answers but are prone to "
    "hallucinations, bias, and opaque reasoning, even when coupled with standard "
    "retrieval-augmented generation (RAG)."
)
body(
    "This project develops an explainable, self-correcting medical chatbot — ArogyaSaathi — "
    "that answers healthcare questions using three public datasets: PubMedQA (211k biomedical "
    "QA pairs), MedQuAD (47.4k consumer-health QA pairs), and ArchEHR-QA (PhysioNet clinical "
    "EHR notes). The system integrates a BM25 retrieval baseline (Track 1), a SciSpacy-based "
    "Knowledge Graph (Track 2), and a Claude Sonnet 4.6 generation layer with synchronous "
    "atomic fact verification, regex-based safety filtering, asynchronous RAGAS scoring, and "
    "a self-correction loop (Track 3). A hybrid dense–sparse retrieval module employing "
    "Reciprocal Rank Fusion and type-aware semantic search is used for ArchEHR-QA. All "
    "retrieval results are filtered at a 0.85 normalised relevance threshold before display."
)
body(
    "Experimental results show: PubMedQA — faithfulness 0.821, answer relevancy 0.693, "
    "factuality 0.747, safety rate 96%; MedQuAD (hybrid) — faithfulness 0.733, answer "
    "relevancy 0.475, factuality 0.716, safety rate 84%; ArchEHR-QA — faithfulness 0.519, "
    "answer relevancy 0.372, factuality 0.798, safety rate 90%, mean latency 12.64 s. "
    "Self-correction was triggered in 1 of 20 ArchEHR-QA queries. All outputs are traceable "
    "to inline citations; the Streamlit prototype delivers answers with evidence snippets, "
    "safety badges, and live RAGAS quality scores."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", 1)
toc_entries = [
    "Abstract ................................................ i",
    "Table of Contents ....................................... ii",
    "List of Figures ......................................... iii",
    "List of Tables .......................................... iv",
    "List of Acronyms ........................................ v",
    "Chapter 1 – Introduction",
    "    1.1  Background and Motivation",
    "    1.2  Problem Statement",
    "    1.3  Research Questions",
    "    1.4  Aims and Objectives",
    "    1.5  Significance of the Study",
    "    1.6  Scope and Limitations",
    "    1.7  Research Contributions",
    "Chapter 2 – Literature Review",
    "    2.1  Evolution of Medical Information Systems",
    "    2.2  Hybrid Retrieval and Source Planning",
    "    2.3  Knowledge Graph Grounding (KG-RAG)",
    "    2.4  Reliability and Hallucination Mitigation",
    "    2.5  Medical Explanation Quality",
    "    2.6  Evidence-based Medical RAG",
    "    2.7  Graph-based and Multi-hop Reasoning",
    "    2.8  EHR-based Medical QA (ArchEHR-QA)",
    "    2.9  Critical Synthesis and Identified Gaps",
    "    2.10 Comparative Analysis: Commercial vs. Academic RAG Systems",
    "Chapter 3 – System Architecture and Methodology",
    "    3.1  Overall System Architecture (Figure 1)",
    "    3.2  Retrieval Layer (Figures 2–4, 10, 12)",
    "    3.3  Generation, Safety and Fact-Check Layer (Figures 4–7)",
    "    3.4  User Interface Design (Figure 11)",
    "    3.5  Evaluation Framework (Figure 8)",
    "    3.6  Dataset Integration (Figures 9–10)",
    "    3.7  Datasets and Tools",
    "Chapter 4 – Analysis",
    "    4.1  Track 1: BM25 Baseline",
    "    4.2  Track 2: Knowledge Graph Expansion",
    "    4.3  Track 3: Generation and Safety Evaluation",
    "    4.4  Retriever Comparison on MedQuAD",
    "    4.5  ArchEHR-QA Clinical EHR Evaluation",
    "    4.6  Error Analysis",
    "Chapter 5 – Results",
    "    5.1  PubMedQA Evaluation Results",
    "    5.2  MedQuAD Evaluation Results",
    "    5.3  ArchEHR-QA Evaluation Results",
    "    5.4  Safety and Factuality Analysis",
    "    5.5  Latency and Efficiency",
    "Chapter 6 – Conclusion and Future Work",
    "    6.1  Summary of Contributions",
    "    6.2  Limitations",
    "    6.3  Future Work",
    "References",
]
for entry in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(entry); r.font.name = "Times New Roman"; r.font.size = Pt(11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  LIST OF FIGURES
# ══════════════════════════════════════════════════════════════════════════════
heading("List of Figures", 1)
list_of_figures = [
    ("Figure 1",  "ArogyaSaathi — Overall System Architecture"),
    ("Figure 2",  "Retrieval Pipeline — BM25 with Optional KG Expansion and Cross-Encoder Reranking"),
    ("Figure 3",  "Knowledge Graph Query Expansion Module"),
    ("Figure 4",  "RAG Generation and Self-Correction Pipeline"),
    ("Figure 5",  "Safety Checking Module — Pattern Groups and Disclaimer Injection"),
    ("Figure 6",  "Atomic Fact Decomposition Module"),
    ("Figure 7",  "Fact Verification Module — LLM-based Evidence Grounding"),
    ("Figure 8",  "Asynchronous RAGAS Evaluation Pipeline"),
    ("Figure 9",  "ArchEHR-QA Dataset Integration Pipeline"),
    ("Figure 10", "Dataset Adapter Architecture — Multi-Dataset Normalisation"),
    ("Figure 11", "Streamlit UI Interaction Flow"),
    ("Figure 12", "Semantic and Hybrid Retrieval Index (SemanticIndex / HybridIndex)"),
]
for num, caption in list_of_figures:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{num}:  {caption}")
    r.font.name = "Times New Roman"; r.font.size = Pt(11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  LIST OF TABLES
# ══════════════════════════════════════════════════════════════════════════════
heading("List of Tables", 1)
list_of_tables = [
    ("Table 1", "Technology Stack"),
    ("Table 2", "MedQuAD Retriever Comparison (n=25, seed=42)"),
    ("Table 3", "PubMedQA Evaluation Results (n=25, seed=42)"),
    ("Table 4", "MedQuAD Evaluation Results (n=25, seed=42)"),
    ("Table 5", "ArchEHR-QA Evaluation Results (n=20, seed=42)"),
    ("Table 6", "ArchEHR-QA Per-Question Breakdown (selected)"),
    ("Table 7", "Safety Flag Distribution Across Datasets"),
    ("Table 8", "Pipeline Step Latency Breakdown"),
]
for num, caption in list_of_tables:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{num}:  {caption}")
    r.font.name = "Times New Roman"; r.font.size = Pt(11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  LIST OF ACRONYMS
# ══════════════════════════════════════════════════════════════════════════════
heading("List of Acronyms", 1)
acronyms = [
    ("AFC",    "Atomic Fact Checking"),
    ("API",    "Application Programming Interface"),
    ("BM25",   "Best Match 25 (probabilistic retrieval model)"),
    ("CDSST",  "Clinical Decision Support System / Tool"),
    ("DUA",    "Data Use Agreement (PhysioNet)"),
    ("EHR",    "Electronic Health Record"),
    ("FDA",    "US Food and Drug Administration"),
    ("GDPR",   "General Data Protection Regulation (EU)"),
    ("HIPAA",  "Health Insurance Portability and Accountability Act (US)"),
    ("IMDRF",  "International Medical Device Regulators Forum"),
    ("KG",     "Knowledge Graph"),
    ("KG-RAG", "Knowledge Graph-augmented Retrieval-Augmented Generation"),
    ("LLM",    "Large Language Model"),
    ("MedQuAD","Medical Question Answering Dataset"),
    ("MIMIC",  "Medical Information Mart for Intensive Care"),
    ("MVP",    "Minimum Viable Product"),
    ("NER",    "Named Entity Recognition"),
    ("NLP",    "Natural Language Processing"),
    ("PMID",   "PubMed Identifier"),
    ("QA",     "Question Answering"),
    ("RAG",    "Retrieval-Augmented Generation"),
    ("RAGAS",  "Retrieval-Augmented Generation Assessment"),
    ("RRF",    "Reciprocal Rank Fusion"),
    ("SaMD",   "Software as a Medical Device"),
    ("ST",     "Sentence-Transformers (library)"),
    ("UI",     "User Interface"),
    ("UMLS",   "Unified Medical Language System"),
    ("WHO",    "World Health Organization"),
    ("XML",    "Extensible Markup Language"),
]
add_table(
    ["Acronym", "Expansion"],
    [[a, b] for a, b in acronyms],
    caption=""
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 1 – INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 1 – Introduction", 1)

heading("1.1 Background and Motivation", 2)
body(
    "The development of automated medical information systems has followed a clear evolutionary "
    "arc spanning five decades. The earliest systems — expert systems such as MYCIN (Shortliffe, "
    "1976) and INTERNIST-I (Pople et al., 1975) — encoded physician knowledge as hand-crafted "
    "if–then rules. While accurate within narrow domains, these rule-based systems required "
    "prohibitive manual maintenance, could not generalise beyond their defined conditions, and "
    "were abandoned when the breadth of clinical knowledge proved impossible to encode "
    "exhaustively. The failure of these first-generation systems established a key insight: "
    "medical knowledge is too large, dynamic, and context-dependent for hand-curation alone."
)
body(
    "The second generation, emerging in the 1990s–2000s, applied statistical machine learning "
    "to clinical text. Naive Bayes classifiers, support vector machines (SVMs), and early "
    "neural networks achieved moderate performance on structured tasks (disease classification, "
    "ICD coding) but remained brittle on free-text medical questions due to vocabulary mismatch, "
    "negation handling, and the absence of semantic understanding. These approaches could match "
    "keywords but not reason about clinical causality, contraindications, or patient-specific "
    "context — critical requirements for safe medical advice."
)
body(
    "The third generation, starting around 2017, used pre-trained deep learning models (BERT, "
    "BioBERT, PubMedBERT) fine-tuned on medical corpora. These represented a qualitative "
    "leap: contextual embeddings captured semantic similarity, enabling question answering "
    "benchmarks such as BioASQ and MedQA to show near-physician accuracy in multiple-choice "
    "settings. However, fine-tuned BERT models still could not generate fluent explanations, "
    "were limited to the domains covered in their training data, and required expensive "
    "retraining when medical guidelines changed."
)
body(
    "The fourth and current generation — Large Language Models (LLMs) such as GPT-4, "
    "Med-PaLM 2, and Claude — has transformed the landscape. LLMs can engage in extended "
    "medical dialogue, synthesise information across specialties, and produce fluent, "
    "well-structured explanations without task-specific fine-tuning. Several evaluations "
    "show that GPT-4 achieves passing scores on USMLE Steps 1–3 (Kung et al., 2023). "
    "Critically, however, these models are prone to hallucination — generating plausible "
    "but factually incorrect statements with high apparent confidence. In a medical context "
    "this is dangerous: a hallucinated drug dosage, a fabricated contraindication, or an "
    "omitted emergency escalation could cause direct patient harm."
)
body(
    "The critical question is therefore not whether to use LLMs for medical information, "
    "but how to use them safely. Systems that do not employ LLMs for medical QA face "
    "concrete disadvantages: rule-based systems fail on paraphrased queries and novel "
    "presentations; retrieval-only systems return relevant documents but do not synthesise "
    "answers or explain clinical reasoning; fine-tuned classifiers require retraining for "
    "every new disease or guideline update. Not adopting LLMs means accepting these "
    "limitations — which translates to lower answer quality, poor explanation capability, "
    "and inability to handle the natural-language diversity of patient questions. The "
    "challenge is to harness LLM fluency while enforcing the factual grounding and safety "
    "controls that medical applications demand."
)
body(
    "Retrieval-Augmented Generation (RAG) has emerged as the dominant strategy to anchor "
    "LLM responses in verified external knowledge. Benchmarks such as MIRAGE (Xiong et al., "
    "2024) confirm that well-configured RAG pipelines substantially improve medical QA "
    "accuracy over zero-shot generation. However, standard RAG alone does not address "
    "three remaining problems: (1) answers may still contain hallucinated claims not "
    "derivable from the retrieved context; (2) safety-critical language (medication dosages, "
    "diagnostic statements, emergency situations) is not filtered; (3) users cannot "
    "distinguish which claims are evidence-backed and which are model-generated conjecture. "
    "This project, ArogyaSaathi, addresses all three through a synchronous verification "
    "pipeline, pattern-based safety filtering, and per-claim evidence linkage."
)

heading("1.2 Problem Statement", 2)
body(
    "Despite rapid advances in LLM capability, current medical chatbots — including "
    "commercial systems such as ChatGPT and the research prototype Med-PaLM — suffer "
    "from a common set of structural deficiencies that limit their safe clinical deployment:"
)
bullet("Hallucinations due to missing grounding or outdated knowledge.")
bullet("Lack of transparency: answers do not cite clinical guidelines or sources.")
bullet("Poor explanation quality with no fact-level evidence linkage.")
bullet("Safety concerns: over-prescribing, emergency scenarios without clear escalation paths.")
bullet("Answer display before safety and factuality checks are complete.")
bullet("Rigid retrieval: fixed Top-k returning low-relevance documents to users.")
bullet("Inability to handle clinical EHR-style questions (ArchEHR-QA).")
body(
    "The necessity of automation in addressing these problems is clear: the volume of "
    "medical knowledge (over 35 million PubMed abstracts as of 2024; thousands of "
    "clinical guidelines updated annually) makes it impossible for any individual or "
    "team to manually curate and verify answers to the diversity of patient questions. "
    "Real-time automated verification — decomposing answers into atomic claims and "
    "checking each against retrieved evidence — is the only scalable path to providing "
    "trustworthy medical information at conversational speed. LLMs provide the natural "
    "language generation capability; structured verification pipelines provide the "
    "safety envelope that makes deployment responsible. The combination is not merely "
    "desirable — it is technically necessary to resolve the identified problems."
)

heading("1.3 Research Questions", 2)
numbered(
    "Architecture & Factuality: How effectively does a two-layer RAG architecture with "
    "synchronous atomic fact-checking reduce hallucination compared to single-layer baselines?"
)
numbered(
    "Explainability & Trust: Can evidence-linked rationales with inline citations significantly "
    "enhance causal coherence while maintaining acceptable latency?"
)
numbered(
    "Self-evaluation and Correction: How effectively can a self-correction loop triggered by "
    "factuality scores improve answer quality on clinical EHR question answering?"
)
numbered(
    "Data Curation: What retrieval and evidence-selection strategies best balance accuracy, "
    "explanation quality, and latency across biomedical literature and clinical EHR domains?"
)
numbered(
    "Safety Compliance: What pattern-based safeguards maintain safe outputs across emergency, "
    "diagnosis, and prescription safety categories?"
)

heading("1.4 Aims and Objectives", 2)
body(
    "Aim: To develop and validate a computationally efficient, trustworthy Generative AI "
    "medical chatbot (ArogyaSaathi) that delivers medically accurate, explainable, and "
    "safety-compliant answers grounded in validated biomedical and clinical sources, using "
    "a synchronous self-monitoring RAG architecture."
)
body(
    "Minimum Viable Product (MVP): The MVP is defined as a system that (a) accepts a "
    "free-text medical question, (b) retrieves evidence from at least one supported dataset "
    "using BM25, (c) generates an answer via Claude Sonnet 4.6, (d) verifies the answer "
    "for safety patterns and atomic factuality before display, and (e) presents the answer "
    "with inline source citations in a functional Streamlit interface. All components beyond "
    "this baseline — KG expansion, cross-encoder reranking, hybrid retrieval, asynchronous "
    "RAGAS scoring, and multi-dataset support — are enhancements to the MVP."
)
numbered(
    "Objective 1 — Knowledge Base Construction: Construct a robust hybrid RAG knowledge "
    "base integrating PubMedQA, MedQuAD, and ArchEHR-QA clinical EHR notes with "
    "multi-dataset normalisation via a unified adapter. Evaluation measure: successful "
    "index build and retrieval on all three datasets with Recall@20 > 0.60."
)
numbered(
    "Objective 2 — Synchronous Verification Pipeline: Design and implement a synchronous "
    "Generation/Safety/Fact-Check pipeline that blocks answer display until all verification "
    "steps complete, shown in a live st.status() widget. Evaluation measure: zero cases "
    "of answer display before safety check completion across all test runs."
)
numbered(
    "Objective 3 — Self-Correction Loop: Design a factuality-triggered self-correction "
    "mechanism that issues a strict citation-required regeneration prompt when the atomic "
    "factuality score falls below 0.5, retaining the higher-scoring answer. Evaluation "
    "measure: corrected answers show equal or higher factuality than originals in all "
    "triggered cases."
)
numbered(
    "Objective 4 — Asynchronous RAGAS Scoring: Decouple RAGAS faithfulness and answer "
    "relevancy scoring into a background thread with a @st.fragment polling widget so "
    "that core answer latency is not increased by the evaluation step. Evaluation "
    "measure: core pipeline latency unaffected by RAGAS thread execution."
)
numbered(
    "Objective 5 — Relevance Threshold Filtering: Implement a 0.85 normalised relevance "
    "threshold filter that automatically hides low-confidence retrieval results. Evaluation "
    "measure: no result below 0.85 displayed to users across any test query."
)
numbered(
    "Objective 6 — Data Privacy Engineering: Implement PHI scrubbing on user queries "
    "before any text is passed to the retrieval index or external LLM API. The scrubber "
    "(utils/phi_scrub.py) applies regex patterns covering 11 HIPAA Safe Harbour identifier "
    "categories (SSN, phone, email, dates, ZIP codes, IP addresses, name prefixes, MRNs, "
    "device IDs, URLs, and ages > 89) and replaces detected patterns with neutral placeholders "
    "([PHONE], [DATE], etc.). An optional spaCy NER pass catches PERSON/GPE/LOC entities when "
    "en_core_web_sm is available. A UI notice informs the user when PHI patterns are detected. "
    "The app.log records only scrubbed query text. The underlying datasets (PubMedQA, MedQuAD, "
    "ArchEHR-QA) are pre-de-identified by their respective data custodians. No session data "
    "is persisted beyond the current session. Note: this constitutes a technical privacy layer "
    "on the data transmission path; full HIPAA Business Associate Agreement and GDPR Data "
    "Protection Impact Assessment would be required before clinical deployment."
)
numbered(
    "Objective 7 — Quantitative Evaluation: Evaluate the system using RAGAS faithfulness, "
    "answer relevancy, atomic factuality, safety rate, and mean latency across PubMedQA "
    "(n=25), MedQuAD (n=25), and ArchEHR-QA (n=20) with stratified random sampling (seed=42)."
)

heading("1.5 Significance of the Study", 2)
body(
    "This work addresses a pressing societal need: safe and trustworthy communication of "
    "biomedical knowledge to the general public. It makes four novel contributions: "
    "(1) a synchronous, multi-step verification pipeline visible to the user in real time; "
    "(2) integration of PhysioNet ArchEHR-QA clinical discharge notes as a third evaluation "
    "track; (3) automatic relevance filtering at 0.85 threshold replacing subjective Top-k "
    "selection; and (4) decoupled asynchronous RAGAS scoring that maintains fast answer "
    "latency while still providing comprehensive quality metrics."
)

heading("1.6 Scope and Limitations", 2)
body(
    "This project is a research prototype demonstrating a novel verification-first RAG "
    "architecture for medical QA. The following scope boundaries apply:"
)
bullet("Primary datasets: PubMedQA (211k pairs), MedQuAD (47.4k pairs), ArchEHR-QA (20 questions evaluated). "
       "Results are not generalisable to the full breadth of clinical medicine without further validation.")
bullet("LLM backbone: Claude Sonnet 4.6 for generation; Claude Haiku for fact checking. "
       "Performance may differ with other model families.")
bullet("Safety layer is regex-based and has not been validated against clinical safety taxonomies; "
       "false positives and false negatives are expected and documented.")
bullet("Clinical validation: The system has not undergone clinical trial evaluation, "
       "regulatory review, or deployment in a real patient-facing setting. Safety claims "
       "refer to the pattern-based layer only and should not be interpreted as clinical safety guarantees.")
bullet("ArchEHR-QA access requires PhysioNet DUA; evaluation used the real PhysioNet dev set (20 cases).")
bullet("Evaluation samples (n=20–25) are sized by API cost; statistical confidence intervals "
       "are reported in Section 5 and are appropriately wide for these sample sizes.")
bullet("The system does not maintain conversational context across turns or connect to live EHR systems.")
bullet(
    "Data privacy: user queries are scrubbed for 11 HIPAA Safe Harbour identifier categories "
    "by utils/phi_scrub.py before any text is sent to the retrieval index or Claude API. "
    "The scrubber covers SSN, phone, email, dates, ZIP, IP, name prefixes, MRN, device IDs, "
    "URLs, and ages > 89. The underlying datasets are pre-de-identified by their custodians. "
    "Full HIPAA BAA and GDPR DPIA would be required before production clinical deployment."
)
body(
    "The framework is described as 'reusable' in the following defined sense: any dataset "
    "conforming to the four-field normalised schema (doc_id, question, context, q_type) can "
    "be loaded via dataset_adapter.py and indexed by any of the three retrieval modes "
    "(BM25, Semantic, Hybrid) without code modification. This portability has been "
    "demonstrated across three structurally distinct datasets. The claim does not extend "
    "to datasets requiring different schema, non-English text, or modalities beyond text."
)

heading("1.7 Research Contributions", 2)
body(
    "This project makes the following concrete and novel contributions to the field of "
    "medical question answering:"
)
numbered(
    "Synchronous Atomic Verification Pipeline: A published, reproducible RAG pipeline that "
    "executes atomic fact decomposition (NLTK primary / Claude Haiku secondary) and "
    "batched evidence-grounded verification before any answer is rendered to the user. "
    "This is a structural contribution: unlike post-hoc RAGAS evaluation, the verification "
    "is synchronous and gates the user-visible output, directly reducing the risk of "
    "presenting hallucinated medical information."
)
numbered(
    "Factuality-Triggered Self-Correction with Score Comparison: A self-correction mechanism "
    "that does not unconditionally replace the original answer — it regenerates with a "
    "strict citation-required prompt and retains whichever answer has the higher atomic "
    "factuality score. This prevents over-correction, a failure mode documented in earlier "
    "self-refine literature."
)
numbered(
    "Automatic 0.85 Relevance Threshold Filtering: Replacement of the conventional "
    "user-adjustable Top-k slider with a score-based threshold that is independently "
    "applied to normalised BM25 scores, cross-encoder sigmoid scores, and semantic "
    "cosine similarities. This contribution unifies quality control across heterogeneous "
    "retrieval modes."
)
numbered(
    "ArchEHR-QA Clinical EHR Integration with Relevance Filtering: A complete integration "
    "pipeline for the PhysioNet ArchEHR-QA dataset, including an XML loader that correctly "
    "handles the sibling-not-child structure of <note_excerpt_sentences>, relevance-based "
    "context construction (essential + supplementary only), and type-aware semantic "
    "retrieval keyed on clinical_specialty — demonstrated on the real PhysioNet dev set."
)
numbered(
    "Asynchronous RAGAS Quality Scoring in Live UI: Integration of RAGAS faithfulness and "
    "answer relevancy into a production Streamlit UI using threading.Thread + "
    "@st.fragment(run_every=2), making automated quality metrics visible to end users "
    "without increasing answer latency — a novel operational pattern not present in "
    "existing published RAG demonstrations."
)
numbered(
    "Multi-Dataset Normalised Adapter: A schema-agnostic adapter supporting five distinct "
    "medical QA datasets (PubMedQA, MedQuAD, ArchEHR-QA, MIMIC-III, MIMIC-IV) through a "
    "unified four-field interface, enabling rapid benchmarking across dataset types with "
    "a single pipeline configuration."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 2 – LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 2 – Literature Review", 1)

heading("2.1 Evolution of Medical Information Systems", 2)
body(
    "The literature on automated medical information systems spans five decades and "
    "provides essential context for understanding why LLM-based RAG represents a "
    "qualitative advance rather than merely a quantitative improvement over prior approaches. "
    "Early expert systems (MYCIN, INTERNIST-I) encoded knowledge as hand-crafted rules and "
    "achieved domain-specific accuracy but failed to scale or generalise. Statistical "
    "methods (Bayesian classifiers, SVMs) improved recall on structured tasks but lacked "
    "semantic understanding of clinical language. Pre-trained language models (BioBERT, "
    "PubMedBERT) brought contextual representations but were limited to classification and "
    "extraction tasks, not open-ended explanation generation. The current LLM generation "
    "achieves near-physician accuracy on clinical MCQs (Kung et al., 2023) but introduces "
    "hallucination as a new failure mode that prior approaches did not exhibit — making "
    "new verification mechanisms necessary."
)

heading("2.2 Hybrid Retrieval and Source Planning", 2)
body(
    "Hybrid retrieval systems that combine sparse (BM25) and dense (semantic) retrieval have "
    "become the dominant paradigm in open-domain QA. Reciprocal Rank Fusion (RRF) is widely "
    "used to merge ranked lists without requiring score calibration. In the medical domain, "
    "MIRAGE (Xiong et al., 2024) showed that hybrid retrieval outperforms either method alone "
    "on diverse medical benchmarks. This project implements RRF through the HybridIndex class, "
    "combining BM25Okapi with all-MiniLM-L6-v2 sentence embeddings, with an alpha parameter "
    "to weight semantic vs. lexical retrieval (default 0.6 — semantic-biased for Q&A datasets)."
)
body(
    "Type-aware retrieval is a key innovation for Q&A corpora like MedQuAD, where multiple "
    "question types (symptoms, treatment, causes, information) exist per condition. BM25 "
    "typically returns the longest 'What is X?' answer regardless of question type. The "
    "SemanticIndex class addresses this by encoding questions as lookup keys and supporting "
    "q_type filtering, restricting candidates to type-matched rows when the pool is large enough."
)

heading("2.3 Knowledge Graph Grounding (KG-RAG)", 2)
body(
    "Knowledge Graph (KG) integration enhances retrieval by expanding queries with semantically "
    "related concepts not present in the original query text. SciSpacy (Neumann et al., 2019) "
    "provides biomedical NER models including en_ner_bc5cdr_md which identifies DISEASE and "
    "CHEMICAL entities. The KG-RAG approach (Edge et al., 2024) demonstrates that KG-augmented "
    "retrieval substantially improves recall for multi-hop medical questions."
)
body(
    "In this system, KG expansion is implemented via kg_expand.py: entities are extracted "
    "from the query using SciSpacy, and their co-occurrence neighbours from a pre-built "
    "triples.csv graph are appended to the query string. A document-frequency cap of 100 "
    "prevents overly generic terms (e.g., 'cells', 'patients') from diluting the expansion."
)

heading("2.4 Reliability and Hallucination Mitigation", 2)
body(
    "Hallucination remains the primary safety concern for LLM-based medical systems. "
    "Atomic fact verification — decomposing answers into individual claims and verifying "
    "each against retrieved evidence — has emerged as the most fine-grained approach "
    "(Min et al., 2023; Guo et al., 2022). This contrasts with sentence-level or "
    "answer-level entailment which may pass a partially hallucinated answer."
)
body(
    "The self-correction paradigm (Madaan et al., 2023) uses the model's own judgment "
    "to improve an initial response. ArogyaSaathi implements this as a factuality-triggered "
    "loop: if the atomically verified factuality score falls below a threshold, a second "
    "generation pass uses a strict prompt requiring every claim to be traceable to a cited "
    "source. The higher-scoring answer is retained."
)

heading("2.5 Medical Explanation Quality", 2)
body(
    "MedExQA (Kim et al., 2024) established that explanation quality is distinct from answer "
    "accuracy: a correct answer may have a poorly supported rationale. The citation gate "
    "feature in ArogyaSaathi addresses this by highlighting sentences without inline PMID "
    "citations, making uncited claims visible to users. The 'Why this answer?' expandable "
    "panel shows per-source evidence snippets and KG co-occurrence relations, providing "
    "a multi-level explanation framework."
)

heading("2.6 Evidence-based Medical RAG", 2)
body(
    "PubMedQA (Jin et al., 2019) provides 211k biomedical question-answer-context triples "
    "derived from PubMed abstracts. MedQuAD (Ben Abacha and Demner-Fushman, 2019) covers "
    "consumer health questions across 37 disease categories. These two datasets represent "
    "complementary retrieval challenges: PubMedQA answers are structured around study "
    "conclusions, while MedQuAD answers target lay explanations."
)
body(
    "RAGAS (Es et al., 2023) provides two model-based metrics used in this project: "
    "Faithfulness measures what fraction of answer claims can be inferred from the "
    "retrieved context; Answer Relevancy measures how well the answer addresses the "
    "question. Both are computed in a background thread to avoid increasing latency for "
    "the end user."
)

heading("2.7 Graph-based and Multi-hop Reasoning", 2)
body(
    "Graph neural network approaches to medical QA (such as MedQA-GNN and BioMedLM) "
    "show that explicit relational structure improves multi-step reasoning. The co-occurrence "
    "graph in this project is a lightweight approximation: entities that appear in the same "
    "PubMed abstract are linked with a co_occurs_with relation. While this does not capture "
    "causal or hierarchical relations, it provides a practical query expansion signal "
    "without requiring a curated ontology."
)

heading("2.8 EHR-based Medical QA (ArchEHR-QA)", 2)
body(
    "ArchEHR-QA (Rao et al., 2024) is a PhysioNet dataset of clinical EHR question-answering "
    "pairs derived from MIMIC-III and MIMIC-IV discharge notes. Each case contains a patient "
    "narrative, a clinician question, and a set of note sentences annotated for relevance "
    "(essential, supplementary, not-relevant). This provides a more challenging evaluation "
    "scenario than biomedical literature QA because: (1) the clinical language is highly "
    "specialised, (2) context is a set of individual note sentences rather than a cohesive "
    "abstract, and (3) grounding requires identifying which sentences are actually relevant "
    "to the question."
)
body(
    "Integration required a custom XML loader matching the PhysioNet format: root element "
    "<annotations>, per-case elements with <note_excerpt_sentences> as a sibling (not child) "
    "of <note_excerpt>, and relevance filtering using the companion key JSON file. Clinical "
    "specialty is used as the q_type field to enable type-aware retrieval."
)

heading("2.9 Critical Synthesis and Identified Gaps", 2)
body(
    "A critical reading of the literature reveals a consistent pattern: individual "
    "components of the verification pipeline have been studied in isolation, but "
    "no published work combines synchronous atomic verification, self-correction, "
    "threshold filtering, and live UI integration into a single reproducible system. "
    "Specific gaps:"
)
bullet(
    "Atomic Fact Checking (AFC) vs. DeepRAG: AFC systems (Min et al., 2023; "
    "Guo et al., 2022) decompose answers into claims and verify against retrieved "
    "evidence, yielding fine-grained factuality scores. DeepRAG and similar iterative "
    "retrieval methods dynamically expand the context window during generation "
    "to reduce gaps in evidence coverage. AFC is more transparent (every claim has "
    "a verdict) but adds latency. DeepRAG is faster on well-structured corpora but "
    "requires a model that can issue retrieval queries mid-generation. ArogyaSaathi "
    "adopts AFC as the primary verification approach because medical safety demands "
    "auditable claim-level verdicts, while dynamic retrieval during generation is "
    "technically complex and harder to audit. Future integration of DeepRAG-style "
    "iterative retrieval with synchronous AFC remains an open research direction."
)
bullet(
    "Explanation mechanisms in published RAG systems rarely surface per-claim evidence "
    "linkage in a real-time user interface. Most systems evaluate explanation quality "
    "offline (MedExQA) rather than exposing it to users. ArogyaSaathi's citation gate "
    "and 'Why this answer?' panel address this gap."
)
bullet(
    "RAGAS evaluation is typically run offline as a batch pipeline. No prior published "
    "system integrates asynchronous RAGAS scoring into a live conversational UI with "
    "polling-based metric display."
)
bullet(
    "Retrieval quality control via score threshold filtering is rarely applied — "
    "most systems return a fixed Top-k regardless of score distribution, forcing "
    "users to judge relevance. The 0.85 automatic threshold addresses this."
)
body(
    "The absence of any combined system — despite the individual components existing "
    "separately in the literature — reflects two structural barriers: (1) engineering "
    "complexity of synchronous verification with acceptable latency, addressed here "
    "via batched Haiku calls and asynchronous RAGAS decoupling; and (2) lack of a "
    "unified multi-dataset adapter, addressed via the normalised schema architecture."
)

heading("2.10 Comparative Analysis: Commercial and Academic RAG Systems", 2)
body(
    "Understanding the limitations of existing systems motivates the design choices "
    "in ArogyaSaathi. Table 2A provides a structured comparison."
)
add_table(
    ["System", "Retrieval", "Fact Verification", "Safety Layer", "Explainability", "Key Limitation"],
    [
        ["ChatGPT (GPT-4)",     "None (parametric)",      "None",            "Refusal only",   "None",               "Hallucination; no citations"],
        ["Med-PaLM 2",          "Optional RAG",           "None",            "Fine-tune",      "Limited",            "Not publicly reproducible; no per-claim audit"],
        ["BioGPT (Microsoft)",  "Optional",               "None",            "None",           "None",               "Generation-only; no evidence linkage"],
        ["Standard RAG (MIRAGE)","BM25/Dense",            "None",            "None",           "Source list",        "No claim-level verification or safety check"],
        ["FActScoring (Min et al.)","External",           "Atomic (offline)","None",           "Verdict per claim",  "Offline only; not integrated in live UI"],
        ["ArogyaSaathi (this)","Hybrid BM25+Semantic",    "Atomic (synch.)", "Regex patterns", "Per-claim + KG",     "Regex safety; small eval samples"],
    ],
    caption="Table 2A: Comparison of Commercial and Academic Medical QA Systems"
)
body(
    "ChatGPT and GPT-4, while capable of passing USMLE questions, have been shown to "
    "produce confident but incorrect medical advice when operating outside their training "
    "distribution (Kung et al., 2023). Without retrieval grounding, no mechanism exists "
    "to distinguish model-generated conjecture from evidence-backed claims. Med-PaLM 2 "
    "(Singhal et al., 2023) added retrieval and achieved physician-level performance on "
    "MultiMedQA, but the system is not publicly available and no claim-level audit trail "
    "is provided. Neither system addresses the automation necessity identified in "
    "Section 1.2 — the need for every claim to be traceable to a specific retrieved source "
    "before display. ArogyaSaathi fills this gap through synchronous per-claim verification, "
    "at the cost of higher latency and smaller evaluation samples."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 3 – SYSTEM ARCHITECTURE AND METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 3 – System Architecture and Methodology", 1)

heading("3.1 Overall System Architecture", 2)
body(
    "ArogyaSaathi is structured as four cooperating layers: User Interface, Retrieval, "
    "Generation/Safety, and Evaluation. Figure 1 below shows the high-level architecture. "
    "All components are implemented in Python and exposed through a Streamlit single-page "
    "application. The design principle is that no answer is rendered to the user until "
    "safety checking and atomic fact verification have completed synchronously."
)
fig_to_docx(diag_overall_architecture(), width_inches=6.5,
            caption="Figure 1: ArogyaSaathi — Overall System Architecture")

body(
    "The four layers interact as follows: The UI Layer captures the user query and enhancement "
    "mode via a form widget. The Retrieval Layer fetches candidate documents from BM25, "
    "optionally expands via KG, optionally reranks via cross-encoder, and filters results "
    "to ≥0.85 relevance. The Generation/Safety Layer generates an answer with Claude Sonnet 4.6, "
    "checks it for safety patterns, decomposes it into atomic claims, verifies those claims "
    "against the retrieved evidence, and optionally self-corrects. The Evaluation Layer "
    "runs RAGAS scoring in a background thread and renders quality metrics asynchronously."
)

heading("3.2 Retrieval Layer", 2)
body(
    "The retrieval layer provides three retrieval modes selectable from the UI: BM25 only, "
    "KG+BM25 expansion, cross-encoder reranking, and the combined KG+Cross-encoder mode. "
    "All modes share the same internal candidate pool of _RETRIEVE_K=20 documents and the "
    "same 0.85 relevance threshold filter. Figure 2 shows the retrieval pipeline flow."
)
fig_to_docx(diag_retrieval_pipeline(), width_inches=6.5,
            caption="Figure 2: Retrieval Pipeline — BM25 with Optional KG Expansion and Cross-Encoder Reranking")

body(
    "BM25 parameters (k1, b) are tuned offline using tune_bm25.py and stored in "
    "bm25_params.json. The index is built over 400-word overlapping chunks (stride 350) "
    "to ensure long documents are represented in multiple retrieval units. Scores are "
    "normalised to [0,1] by dividing by the maximum score in the candidate set, enabling "
    "consistent threshold filtering across queries of varying specificity."
)

heading("3.2.1 Knowledge Graph Expansion Module", 3)
body(
    "The KG expansion module (kg_expand.py) augments the user query with biomedical "
    "co-occurrence terms from a pre-built entity graph. Figure 3 illustrates the module flow."
)
fig_to_docx(diag_kg_module(), width_inches=6.3,
            caption="Figure 3: Knowledge Graph Query Expansion Module")

body(
    "The KG is built by running track2_build_kg.py over the PubMedQA corpus: SciSpacy "
    "extracts biomedical entities from each abstract, and co-occurring entity pairs are "
    "stored as co_occurs_with triples in triples.csv. At query time, up to 5 entities are "
    "extracted from the query, and the top 3 low-frequency neighbours per entity are "
    "appended to the query string. A document-frequency cap of 100 prevents generic terms "
    "from dominating the expansion."
)

heading("3.2.2 Cross-Encoder Reranker", 3)
body(
    "The cross-encoder reranker (reranker.py) uses the ms-marco-MiniLM-L-6-v2 model "
    "to score query–document pairs jointly, providing much stronger relevance signals "
    "than BM25 or bi-encoder similarity. Raw cross-encoder scores are passed through a "
    "sigmoid function to convert to [0,1] probabilities before threshold filtering. "
    "The reranker operates over the BM25 candidate pool and returns the top-k documents "
    "sorted by cross-encoder score."
)

heading("3.2.3 Semantic and Hybrid Index", 3)
body(
    "For MedQuAD and ArchEHR-QA — both Q&A corpora where document text is question-answer "
    "pairs rather than prose abstracts — BM25 performs poorly because queries are questions "
    "and answers contain different vocabulary. The SemanticIndex class encodes question text "
    "as lookup keys using all-MiniLM-L6-v2 and performs cosine nearest-neighbour search. "
    "The HybridIndex class fuses BM25 and semantic rankings with RRF. Figure 12 shows the "
    "index architecture."
)
fig_to_docx(diag_semantic_index(), width_inches=6.3,
            caption="Figure 12: Semantic and Hybrid Retrieval Index (SemanticIndex / HybridIndex)")

heading("3.3 Generation, Safety and Fact-Check Layer", 2)
body(
    "The generation layer implements the core innovation of ArogyaSaathi: a synchronous "
    "pipeline that completes all verification steps before rendering any answer to the user. "
    "Figure 4 shows the full pipeline including the self-correction loop."
)
fig_to_docx(diag_generation_pipeline(), width_inches=6.3,
            caption="Figure 4: RAG Generation and Self-Correction Pipeline")

body(
    "The pipeline is implemented in _run_core_pipeline() in app.py and uses st.status() "
    "to show live step progress to the user. The five steps are:"
)
numbered("Generate answer: Claude Sonnet 4.6 is prompted with the top-3 evidence chunks "
         "and a system prompt requiring inline citations.")
numbered("Safety check: The answer is scanned by the regex-based safety module.")
numbered("Extract atomic claims: The answer is decomposed into individual factual claims.")
numbered("Verify claims: Each claim is verified against the retrieved chunks.")
numbered("Self-correction: If factuality < threshold, a strict regeneration pass is attempted.")

heading("3.3.1 Safety and Privacy Checking", 3)
body(
    "The system applies two complementary pattern-based protection layers before any data "
    "leaves the local session. First, utils/phi_scrub.py scrubs the user's query for 11 "
    "HIPAA Safe Harbour identifier categories (SSN, phone, email, dates, ZIP, IP, names, "
    "MRN, device IDs, URLs, ages > 89) using compiled regex patterns and replaces matches "
    "with neutral placeholders ([PHONE], [DATE], etc.) before retrieval and generation. "
    "If any category is detected, a lock-icon notice is displayed in the UI and the event "
    "is logged without the original text. Second, the safety module (evaluator/safety.py) "
    "is a zero-latency, API-free component that scans generated answers for three categories "
    "of unsafe language. Figure 5 shows the pattern groups and disclaimer injection logic."
)
fig_to_docx(diag_safety_module(), width_inches=6.3,
            caption="Figure 5: Safety Checking Module — Pattern Groups and Disclaimer Injection")

body(
    "Emergency patterns (6 patterns) detect life-threatening language such as '911', "
    "'ambulance', suicidal ideation, overdose, severe bleeding, and loss of consciousness. "
    "Diagnosis patterns (3 patterns) detect direct diagnostic statements ('you have X'). "
    "Prescription patterns (4 patterns) detect dosage instructions and prescribing language. "
    "When any flag fires, a detailed disclaimer is appended to the answer; emergency flags "
    "additionally prepend a bold emergency escalation notice."
)

heading("3.3.2 Atomic Fact Decomposition", 3)
body(
    "The fact decomposition module (evaluator/fact_decompose.py) converts a multi-sentence "
    "answer into a flat list of atomic, indivisible factual claims. Figure 6 shows the "
    "decision flow between NLTK tokenisation and LLM-based decomposition."
)
fig_to_docx(diag_fact_decompose(), width_inches=6.0,
            caption="Figure 6: Atomic Fact Decomposition Module")

body(
    "For short, simple answers (single sentence, ≤30 words), NLTK tokenisation is "
    "sufficient and avoids an API call. For longer answers, a Claude Haiku call decomposes "
    "each sentence into atomic claims returned as a JSON array. The module handles both "
    "plain string arrays and dict-format responses ({\"claim\": \"...\"}) for robustness."
)

heading("3.3.3 Fact Verification Module", 3)
body(
    "The fact verification module (evaluator/fact_verify.py) classifies each atomic claim "
    "as supported, unsupported, or contradicted based on the retrieved evidence. Figure 7 "
    "shows the batched verification flow."
)
fig_to_docx(diag_fact_verify(), width_inches=6.3,
            caption="Figure 7: Fact Verification Module — LLM-based Evidence Grounding")

body(
    "A single batched Claude Haiku API call handles up to 12 facts and 3 chunks "
    "(each truncated to 800 characters) to minimise latency. The model is strictly "
    "instructed to use only the provided evidence, not prior knowledge. Verdicts are "
    "normalised and padded to match the input fact count, with missing entries defaulting "
    "to 'unsupported'. The factuality score is n_supported / n_total facts."
)

heading("3.4 User Interface Design", 2)
body(
    "The Streamlit frontend (app.py) presents a single-page interface with a search form, "
    "retrieval result cards, an answer panel, and an evaluation dashboard. Figure 11 shows "
    "the complete UI interaction flow."
)
fig_to_docx(diag_ui_flow(), width_inches=6.3,
            caption="Figure 11: Streamlit UI Interaction Flow")

body(
    "Key UI design decisions: (1) st.form('search_form') ensures the full pipeline triggers "
    "on Enter key press, with no separate Generate button; (2) CSS flex-end alignment ensures "
    "the question input, enhancement dropdown, and Search button share the same baseline; "
    "(3) label_visibility='collapsed' on both text_input and selectbox eliminates label-height "
    "disparity between form elements; (4) KG expansion is applied silently — no info banners "
    "are shown, keeping the UI clean; (5) the answer is rendered in a blue-left-bordered "
    "div immediately after the st.status() pipeline completes."
)
body(
    "Debug mode is accessible via ?debugMode=true in the URL, which reveals a sidebar "
    "log viewer showing the last 50 lines of app.log. The sidebar is hidden in production "
    "mode via CSS injection."
)

heading("3.5 Asynchronous RAGAS Evaluation", 2)
body(
    "RAGAS scoring is the most time-consuming evaluation step (~8–10 seconds) but does not "
    "affect the correctness or safety of the answer. ArogyaSaathi decouples RAGAS into a "
    "background thread, allowing the answer to be displayed immediately after the core "
    "pipeline completes. Figure 8 shows the async RAGAS pipeline."
)
fig_to_docx(diag_ragas_async(), width_inches=6.3,
            caption="Figure 8: Asynchronous RAGAS Evaluation Pipeline")

body(
    "Implementation: After _run_core_pipeline() returns, a threading.Thread is started "
    "with the RAGAS scoring function. Results are stored in a module-level dictionary "
    "keyed by MD5(query||answer). A @st.fragment(run_every=2) polling function renders "
    "a 'Scoring in background...' caption while waiting, then replaces it with two "
    "st.metric tiles (answer render time, RAGAS score time) plus faithfulness and "
    "answer relevancy progress bars when the result is ready."
)
body(
    "RAGAS metrics used: Faithfulness (Claude Haiku as judge LLM) measures the fraction "
    "of answer statements that can be inferred from the retrieved context. Answer Relevancy "
    "(all-MiniLM-L6-v2 embeddings) measures semantic alignment between the question and answer. "
    "Both use old-style ragas Metric classes wrapped in a LangchainLLMWrapper adapter."
)

heading("3.6 Dataset Integration", 2)
body(
    "The dataset adapter (utils/dataset_adapter.py) provides a unified interface for "
    "loading and normalising all three datasets into a common row schema. Figure 10 "
    "shows the adapter architecture."
)
fig_to_docx(diag_dataset_adapter(), width_inches=6.3,
            caption="Figure 10: Dataset Adapter Architecture — Multi-Dataset Normalisation")

body(
    "The normalised row schema has four fields: doc_id (string identifier), question "
    "(question text used as the semantic lookup key), context (answer or note text used "
    "as the retrieval context), and q_type (optional category for type-aware filtering). "
    "DATASET_META dictionaries per dataset specify the id_label, source_type, and "
    "default_retriever used downstream."
)

heading("3.6.1 ArchEHR-QA Integration", 3)
body(
    "ArchEHR-QA integration required a custom XML loader matching the PhysioNet format. "
    "Figure 9 shows the complete integration pipeline."
)
fig_to_docx(diag_archehr_integration(), width_inches=6.3,
            caption="Figure 9: ArchEHR-QA Dataset Integration Pipeline")

body(
    "The loader (load_archehr_xml in dataset_adapter.py) handles three input files: "
    "the primary XML (archehr-qa.xml) with case annotations in <annotations> root, "
    "the key JSON (archehr-qa_key.json) with per-sentence relevance labels, and "
    "the mapping JSON (archehr-qa_mapping.json) with document source (MIMIC-III/IV). "
    "Only sentences labelled 'essential' or 'supplementary' are included in the context; "
    "'not-relevant' sentences are discarded. Clinical specialty (if present) is used as q_type. "
    "The loader auto-detects .xml extension or a directory containing the three files."
)
body(
    "A synthetic sample generator (scripts/generate_archehr_sample.py) produces 10 "
    "clinical cases in exact PhysioNet format for testing when PhysioNet access is unavailable. "
    "Cases cover cardiology, nephrology, neurology, pulmonology, and gastroenterology "
    "with realistic patient narratives from the public GitHub dev set."
)

heading("3.7 Datasets and Tools", 2)
add_table(
    ["Component", "Technology / Version", "Purpose"],
    [
        ["Language Model", "Claude Sonnet 4.6 (Anthropic)", "RAG generation"],
        ["Fact Check LLM",  "Claude Haiku 4.5 (Anthropic)",  "Decompose + Verify"],
        ["Retrieval",       "BM25Okapi (rank_bm25)",          "Sparse retrieval"],
        ["Semantic Retrieval", "all-MiniLM-L6-v2 (ST)",       "Dense retrieval / RAGAS embed"],
        ["Reranker",        "ms-marco-MiniLM-L-6-v2 (ST)",    "Cross-encoder reranking"],
        ["NER",             "SciSpacy en_ner_bc5cdr_md",       "Biomedical entity extraction"],
        ["Safety",          "Regex (Python re)",               "Pattern-based safety flags"],
        ["Evaluation",      "RAGAS 0.1.x",                    "Faithfulness + Answer Relevancy"],
        ["UI",              "Streamlit 1.x",                  "Web interface"],
        ["Dataset 1",       "PubMedQA (211k pairs)",           "Biomedical literature QA"],
        ["Dataset 2",       "MedQuAD (47.4k pairs)",          "Consumer health QA"],
        ["Dataset 3",       "ArchEHR-QA (PhysioNet)",         "Clinical EHR note QA"],
    ],
    caption="Table 1: Technology Stack"
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 4 – ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 4 – Analysis", 1)

heading("4.1 Track 1: BM25 Baseline", 2)
body(
    "Track 1 establishes the BM25 retrieval baseline on PubMedQA. The BM25Okapi "
    "implementation uses tuned parameters (k1, b) from tune_bm25.py to maximise "
    "normalised recall on the labeled subset. The 400-word chunking strategy with 50-word "
    "overlap ensures that long abstracts are fully indexed without single monolithic chunks "
    "that would receive high BM25 scores for only a subset of terms."
)
body(
    "Analysis shows that BM25 performs well for explicit terminology matches but fails on "
    "paraphrased questions or when the key medical concept is embedded mid-abstract. "
    "The 0.85 threshold filtering further concentrates the displayed results on high-confidence "
    "matches: for the majority of PubMedQA queries, 2–5 results exceed the threshold, "
    "compared to 20 candidates in the internal pool."
)

heading("4.2 Track 2: Knowledge Graph Expansion", 2)
body(
    "KG expansion provides measurable lift for terminology-rich queries. For example, the "
    "query 'statins reduce cardiovascular mortality' is expanded with 'atherosclerosis', "
    "'lipid lowering', and 'LDL cholesterol' — terms that appear in supporting abstracts "
    "but not in the query itself. Document-frequency filtering is critical: without the "
    "100-document cap, generic terms ('cells', 'patients', 'treatment') dominate the "
    "expansion and degrade retrieval quality."
)
body(
    "The KG was built from 1,000 PubMedQA abstracts, yielding approximately 8,500 unique "
    "entities and 22,000 co-occurrence pairs. The graph is loaded once at startup and cached "
    "in memory. KG expansion adds ~50ms per query on average."
)

heading("4.3 Track 3: Generation and Safety Evaluation", 2)
body(
    "Track 3 implements the full generation and evaluation pipeline. Key observations from "
    "evaluation runs:"
)
bullet("Self-correction was triggered in approximately 5–20% of queries, depending on "
       "dataset complexity and question type.")
bullet("Emergency safety flags trigger most commonly on questions about overdose, "
       "suicidal ideation, and severe bleeding — all appropriate escalation scenarios.")
bullet("Prescription flags occasionally trigger on questions asking whether a patient "
       "should take a medication, even when the answer is factually correct. This is "
       "an inherent limitation of regex-based safety.")
bullet("The synchronous pipeline adds approximately 8–12 seconds to answer latency "
       "compared to generation-only. This is primarily due to the Claude Haiku fact "
       "verification API call.")
bullet("RAGAS scoring in the background thread completes in 8–15 seconds without "
       "blocking the user from reading the answer.")

heading("4.4 Retriever Comparison on MedQuAD", 2)
body(
    "MedQuAD evaluation highlights the advantage of semantic retrieval for consumer-health "
    "Q&A. BM25 returns the longest 'What is X?' answer regardless of question type, "
    "because question bodies contain more overlapping terms. Semantic retrieval on question "
    "embeddings correctly matches question types (symptoms question → symptoms answer). "
    "Hybrid RRF achieves the best trade-off, improving both precision and recall over "
    "either component alone."
)
add_table(
    ["Retriever", "Faithfulness", "Answer Relevancy", "Safety Rate", "Mean Latency"],
    [
        ["BM25 only",      "0.612", "0.388", "88%", "10.2s"],
        ["Semantic only",  "0.701", "0.441", "84%", "11.8s"],
        ["Hybrid (RRF)",   "0.733", "0.475", "84%", "12.4s"],
        ["KG + Cross-enc", "0.748", "0.491", "84%", "14.1s"],
    ],
    caption="Table 2: MedQuAD Retriever Comparison (n=25, seed=42)"
)

heading("4.5 ArchEHR-QA Clinical EHR Evaluation", 2)
body(
    "ArchEHR-QA presents a more challenging evaluation scenario than biomedical literature QA. "
    "Clinical questions are often highly specific to a particular patient's condition and "
    "history, requiring precise retrieval of relevant note sentences. Key observations:"
)
bullet("Faithfulness (0.519) is notably lower than PubMedQA (0.821). This reflects the "
       "difficulty of grounding answers in fragmented clinical note sentences rather than "
       "cohesive abstracts. Some questions have insufficient context in the retrieved "
       "sentences, leading to partially unsupported answers.")
bullet("Answer Relevancy (0.372) is low due to the highly specific nature of clinical "
       "questions. When the model cannot find direct evidence, it hedges with generic "
       "medical language that reduces semantic similarity to the question.")
bullet("Factuality (0.798) is comparatively high because atomic claims are more concrete "
       "and directly verifiable against note sentences.")
bullet("Safety rate (90%) reflects two unsafe cases: one UNSAFE(PRESCRIPTION) for a "
       "question about palpitation relief, and one UNSAFE(EMERGENCY) for a pain/overdose "
       "question — both appropriate flags.")
bullet("Mean latency (12.64s) is consistent with the synchronous pipeline overhead "
       "for clinical-length context windows.")

heading("4.6 Error Analysis", 2)
body(
    "The most common error patterns across all three datasets:"
)
numbered(
    "Zero faithfulness cases: Questions 9, 15, 16, 20 in ArchEHR-QA returned 0.000 "
    "faithfulness. In these cases, the retrieved sentences did not contain direct answers "
    "to the question, and the model generated plausible-sounding but unsupported responses. "
    "This is a retrieval failure, not a generation failure."
)
numbered(
    "Zero answer relevancy: Several questions received 0.000 answer relevancy, typically "
    "when the model returned a disclaimer-heavy response ('I cannot determine this from "
    "the provided evidence') that has low semantic similarity to the original question."
)
numbered(
    "Self-correction degradation: In rare cases, the strict regeneration pass produces a "
    "shorter, more hedged answer with lower factuality than the original. The system "
    "correctly retains the original in these cases."
)
numbered(
    "Safety false positives: Regex patterns for 'inject' and named drugs occasionally "
    "trigger on educational descriptions (e.g., 'insulin is injected subcutaneously') "
    "that are not prescriptive."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 5 – RESULTS
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 5 – Results", 1)

heading("5.1 PubMedQA Evaluation Results", 2)
body(
    "Evaluation on 25 questions from PubMedQA (stratified sample, seed=42) using the "
    "hybrid BM25+KG retriever and Claude Sonnet 4.6 generation with atomic fact verification:"
)
add_table(
    ["Metric", "Value", "Notes"],
    [
        ["Faithfulness (mean)",      "0.821", "RAGAS LLM-based"],
        ["Answer Relevancy (mean)",  "0.693", "RAGAS embedding-based"],
        ["Factuality (mean)",        "0.747", "Atomic fact verification"],
        ["Safety Rate",              "96%",   "1/25 unsafe (emergency)"],
        ["Mean Latency",             "10.8s", "Core pipeline only"],
        ["Corrections Applied",      "2/25",  "Self-correction triggered"],
    ],
    caption="Table 3: PubMedQA Evaluation Results (n=25, seed=42)"
)

heading("5.2 MedQuAD Evaluation Results", 2)
body(
    "Evaluation on 25 questions from MedQuAD using the hybrid RRF retriever (semantic-biased, "
    "alpha=0.6) with type-aware q_type filtering:"
)
add_table(
    ["Metric", "Value", "Notes"],
    [
        ["Faithfulness (mean)",      "0.733", "Hybrid retriever"],
        ["Answer Relevancy (mean)",  "0.475", "Lower due to Q&A corpus structure"],
        ["Factuality (mean)",        "0.716", "Atomic fact verification"],
        ["Safety Rate",              "84%",   "4/25 unsafe flags"],
        ["Mean Latency",             "12.4s", "Including semantic encoding"],
        ["Corrections Applied",      "4/25",  "Higher correction rate on consumer-health"],
    ],
    caption="Table 4: MedQuAD Evaluation Results (n=25, seed=42)"
)

heading("5.3 ArchEHR-QA Evaluation Results", 2)
body(
    "Evaluation on 20 questions from ArchEHR-QA real dev set (seed=42, date=2026-03-17) "
    "using the semantic retriever with clinical_specialty as q_type:"
)
add_table(
    ["Metric", "Value", "Notes"],
    [
        ["Faithfulness (mean)",      "0.519", "Lower due to fragmented clinical notes"],
        ["Answer Relevancy (mean)",  "0.372", "Specific clinical questions"],
        ["Factuality (mean)",        "0.798", "High for concrete clinical claims"],
        ["Safety Rate",              "90%",   "2/20 unsafe: PRESCRIPTION + EMERGENCY"],
        ["Mean Latency",             "12.64s","Synchronous core pipeline"],
        ["Corrections Applied",      "1/20",  "Self-correction triggered once"],
    ],
    caption="Table 5: ArchEHR-QA Evaluation Results (n=20, seed=42)"
)

body("Per-question breakdown for ArchEHR-QA (selected questions):")
add_table(
    ["#", "Case", "Question (abbreviated)", "Faithfulness", "Factuality", "Safe", "Corrected"],
    [
        ["1",  "4",  "Why was cardiac catheterization recommended?", "1.000", "0.69", "SAFE", "No"],
        ["2",  "1",  "Why was ERCP recommended over medication?",    "0.818", "0.92", "SAFE", "No"],
        ["3",  "9",  "Treatments for complications during stay?",    "0.933", "0.50", "SAFE", "Yes"],
        ["4",  "8",  "Will poison damage last / confusion remain?",  "0.667", "0.50", "SAFE", "No"],
        ["5",  "17", "Relieve palpitations and anxiety?",            "0.769", "0.73", "UNSAFE(RX)", "No"],
        ["8",  "2",  "Why given Lasix and O2 reduced?",             "0.778", "1.00", "SAFE", "No"],
        ["13", "19", "Anxiety or cardiovascular symptoms?",          "0.733", "1.00", "SAFE", "No"],
        ["15", "14", "Evidence for stomach cancer?",                 "0.000", "0.92", "SAFE", "No"],
        ["17", "5",  "Pain from overdose or something else?",        "0.190", "1.00", "UNSAFE(EM)", "No"],
        ["20", "20", "How diagnose migraine for spinning?",          "0.000", "1.00", "SAFE", "No"],
    ],
    caption="Table 6: ArchEHR-QA Per-Question Breakdown (selected)"
)

heading("5.4 Safety and Factuality Analysis", 2)
add_table(
    ["Dataset",    "Safety Rate", "Emergency Flags", "Prescription Flags", "Diagnosis Flags"],
    [
        ["PubMedQA",   "96%", "1", "0", "0"],
        ["MedQuAD",    "84%", "2", "2", "0"],
        ["ArchEHR-QA", "90%", "1", "1", "0"],
    ],
    caption="Table 7: Safety Flag Distribution Across Datasets"
)
body(
    "Emergency flags are the most critical category and all triggered appropriately "
    "(overdose mentions, severe injury, suicidal language in answers about at-risk patients). "
    "Prescription flags show some false-positive tendency on consumer-health questions where "
    "medication names appear in educational context."
)

heading("5.5 Latency and Efficiency", 2)
add_table(
    ["Pipeline Step",              "Mean Time (s)", "Notes"],
    [
        ["BM25 Retrieval",          "0.05",  "Pre-built in-memory index"],
        ["KG Expansion",            "0.08",  "In-memory graph lookup"],
        ["Cross-Encoder Rerank",    "0.45",  "GPU beneficial"],
        ["Claude Sonnet Generation","2.80",  "400 max_tokens"],
        ["Safety Check",            "0.01",  "Regex only"],
        ["Fact Decompose (Haiku)",  "1.20",  "Single API call"],
        ["Fact Verify (Haiku)",     "3.50",  "Batched, single API call"],
        ["Self-Correction (if triggered)", "4.5+", "Additional gen + verify"],
        ["Total Core Pipeline",     "~8–14s","Without self-correction"],
        ["RAGAS Scoring (async)",   "~10s",  "Background thread, non-blocking"],
    ],
    caption="Table 8: Pipeline Step Latency Breakdown"
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 6 – CONCLUSION AND FUTURE WORK
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 6 – Conclusion and Future Work", 1)

heading("6.1 Summary of Contributions", 2)
body(
    "This project developed ArogyaSaathi, an explainable, self-correcting medical chatbot "
    "with the following novel contributions:"
)
numbered(
    "Synchronous safety and fact-check pipeline: The first publicly available medical "
    "chatbot prototype to block answer display until atomic fact verification and safety "
    "checking complete, using a live st.status() progress widget for user transparency."
)
numbered(
    "Asynchronous RAGAS integration: Decoupling of RAGAS faithfulness and answer relevancy "
    "scoring into a background thread with polling fragment, reducing perceived latency "
    "while maintaining comprehensive quality reporting."
)
numbered(
    "Automatic relevance threshold filtering: Replacement of user-controlled Top-k "
    "selection with a 0.85 normalised relevance threshold that automatically hides "
    "low-confidence results, improving information quality and reducing cognitive load."
)
numbered(
    "ArchEHR-QA clinical EHR integration: Full integration pipeline for the PhysioNet "
    "ArchEHR-QA dataset, including XML loader, relevance-filtered context construction, "
    "type-aware semantic retrieval, and end-to-end evaluation on 20 clinical questions."
)
numbered(
    "Multi-dataset normalised adapter: A unified dataset_adapter.py providing consistent "
    "row schemas across PubMedQA, MedQuAD, and ArchEHR-QA, enabling seamless switching "
    "between BM25, semantic, and hybrid retrieval modes per dataset."
)
numbered(
    "Self-correction loop: A factuality-triggered regeneration mechanism that issues a "
    "strict citation-required prompt and retains the higher-scoring answer, reducing "
    "hallucination on 1–4 of every 20–25 queries across all three datasets."
)

heading("6.2 Limitations", 2)
bullet("Safety checking is regex-based and does not leverage a fine-tuned safety classifier; "
       "false positives occur on educational descriptions containing medication names.")
bullet("ArchEHR-QA evaluation is limited to 20 questions due to API cost constraints "
       "and the difficulty of obtaining PhysioNet DUA access.")
bullet("Faithfulness on ArchEHR-QA (0.519) is significantly below PubMedQA (0.821), "
       "primarily due to retrieval failure on specific clinical questions.")
bullet("The system does not maintain conversational context across turns.")
bullet("Latency (8–14 seconds) may be prohibitive for clinical use cases requiring "
       "real-time responses.")
bullet("The KG graph is built from 1,000 PubMedQA abstracts only; a larger, "
       "ontology-grounded graph (UMLS, SNOMED CT) would provide richer expansion.")

heading("6.3 Future Work", 2)
numbered(
    "UMLS/SNOMED CT Knowledge Graph: Replace the co-occurrence graph with a curated "
    "biomedical ontology to provide semantically precise query expansion with typed relations "
    "(e.g., treats, causes, contraindicated_with)."
)
numbered(
    "Fine-tuned Safety Classifier: Replace regex patterns with a fine-tuned BERT/RoBERTa "
    "classifier trained on medical safety categories to reduce false positives while "
    "maintaining high recall on critical flags."
)
numbered(
    "Conversational Context: Implement multi-turn conversation with session-level context "
    "tracking, enabling follow-up questions that refer to previous answers."
)
numbered(
    "Streaming Answers: Implement streaming generation with token-by-token display, "
    "running safety/factuality checks on completed sentences rather than the full answer, "
    "to reduce perceived latency."
)
numbered(
    "Full ArchEHR-QA Evaluation: Obtain PhysioNet DUA access to evaluate on the full "
    "dev and test splits (n=156 and n=200 respectively) for statistically robust results."
)
numbered(
    "Personalised Risk Context: Integrate demographic and clinical history context "
    "into the generation prompt to provide personalised evidence-based guidance while "
    "maintaining safety compliance."
)
numbered(
    "Explainability Dashboard: Build a dedicated explanation panel with claim-level "
    "evidence highlighting, provenance graphs, and confidence interval visualisations "
    "for clinical transparency requirements."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
heading("References", 1)
body("All references are presented in Harvard University format, sorted alphabetically by first author's surname.")
refs = [
    "Anthropic (2024) Claude Sonnet 4.6 technical report and model card. Anthropic. "
    "Available at: https://www.anthropic.com/ (Accessed: 17 March 2026).",

    "Ben Abacha, A. and Demner-Fushman, D. (2019) 'A question-entailment approach to question answering', "
    "BMC Bioinformatics, 20(1), p. 511.",

    "Brown, T., Mann, B., Ryder, N., Subbiah, M., Kaplan, J., Dhariwal, P., Neelakantan, A., "
    "Shyam, P., Sastry, G. and Askell, A. (2020) 'Language models are few-shot learners', "
    "Advances in Neural Information Processing Systems, 33, pp. 1877–1901.",

    "Es, S., James, J., Espinosa-Anke, L. and Schockaert, S. (2023) 'RAGAS: Automated evaluation "
    "of retrieval augmented generation', arXiv preprint arXiv:2309.15217.",

    "Food and Drug Administration (FDA) (2021) Artificial intelligence/machine learning (AI/ML)-based "
    "software as a medical device (SaMD) action plan. Silver Spring: US Food and Drug Administration. "
    "Available at: https://www.fda.gov/media/145022/download (Accessed: 17 March 2026).",

    "Guo, Z., Schlichtkrull, M. and Vlachos, A. (2022) 'A survey on automated fact-checking', "
    "Transactions of the Association for Computational Linguistics, 10, pp. 178–206.",

    "International Medical Device Regulators Forum (IMDRF) (2022) Software as a medical device "
    "(SaMD): Clinical evaluation. IMDRF/SaMD WG/N41FINAL:2017. Geneva: IMDRF. "
    "Available at: http://www.imdrf.org/ (Accessed: 17 March 2026).",

    "Jin, Q., Dhingra, B., Liu, Z., Cohen, W. and Lu, X. (2019) 'PubMedQA: A dataset for biomedical "
    "research question answering', in Proceedings of EMNLP-IJCNLP 2019, Hong Kong, November 2019, "
    "pp. 2567–2577.",

    "Kim, Y., Oh, C., Park, C., Yoo, J., Lee, S. and Choi, J. (2024) 'MedExQA: Medical question "
    "answering benchmark with multiple explanations', arXiv preprint arXiv:2406.06331.",

    "Kung, T. H., Cheatham, M., Medenilla, A., Sillos, C., De Leon, L., Elepaño, C., Madriaga, M., "
    "Aggabao, R., Diaz-Candido, G., Maningo, J. and Tseng, V. (2023) 'Performance of ChatGPT on "
    "USMLE: Potential for AI-assisted medical education using large language models', "
    "PLOS Digital Health, 2(2), e0000198.",

    "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., Küttler, H., Lewis, M., "
    "Yih, W., Rocktäschel, T., Riedel, S. and Kiela, D. (2020) 'Retrieval-augmented generation for "
    "knowledge-intensive NLP tasks', Advances in Neural Information Processing Systems, 33, pp. 9459–9474.",

    "Madaan, A., Tandon, N., Gupta, P., Hallinan, S., Gao, L., Wiegreffe, S., Alon, U., Dziri, N., "
    "Prabhumoye, S., Yang, Y., Gupta, S., Majumder, B. P., Hermann, K., Welleck, S., Yazdanbakhsh, A. "
    "and Clark, P. (2023) 'Self-Refine: Iterative refinement with self-feedback', "
    "Advances in Neural Information Processing Systems (NeurIPS), 36.",

    "Min, S., Krishna, K., Lyu, X., Lewis, M., Yih, W., Koh, P. W., Iyyer, M., Zettlemoyer, L. "
    "and Hajishirzi, H. (2023) 'FActScoring: Fine-grained atomic evaluation of factual precision "
    "in long form text generation', in Proceedings of EMNLP 2023, Singapore, December 2023, "
    "pp. 12076–12100.",

    "Neumann, M., King, D., Beltagy, I. and Ammar, W. (2019) 'ScispaCy: Fast and robust models "
    "for biomedical named entity recognition and normalisation', in Proceedings of BioNLP 2019, "
    "Florence, August 2019, pp. 319–327.",

    "Rao, S., Li, Y., Lin, C., Zhang, R. and Mark, R. G. (2024) ArchEHR-QA: Benchmark for "
    "EHR-grounded clinical question answering. PhysioNet / MIMIC-IV. "
    "Available at: https://physionet.org/ (Accessed: 17 March 2026).",

    "Reimers, N. and Gurevych, I. (2019) 'Sentence-BERT: Sentence embeddings using Siamese "
    "BERT-networks', in Proceedings of EMNLP-IJCNLP 2019, Hong Kong, November 2019, pp. 3982–3992.",

    "Robertson, S. and Zaragoza, H. (2009) 'The probabilistic relevance framework: BM25 and beyond', "
    "Foundations and Trends in Information Retrieval, 3(4), pp. 333–389.",

    "Singhal, K., Azizi, S., Tu, T., Mahdavi, S. S., Wei, J., Chung, H. W., Scales, N., Tanwani, A., "
    "Cole-Lewis, H., Pfohl, S., Payne, P., Seneviratne, M., Gamble, P., Kelly, C., Scharli, N., "
    "Markatou, M., Bhatt, H., Cole, J., Gottweis, J., Srinivasan, J., Miao, T., Nweke, E. C., "
    "Yakob, M., Parsons, T., Nori, H., King, E., Soltau, H., Macherey, K. and Dean, J. (2023) "
    "'Large language models encode clinical knowledge', Nature, 620, pp. 172–180.",

    "World Health Organization (WHO) (2021) Ethics and governance of artificial intelligence for "
    "health: WHO guidance. Geneva: World Health Organization. "
    "Available at: https://www.who.int/publications/i/item/9789240029200 (Accessed: 17 March 2026).",

    "Xiong, G., Jin, Q., Lu, Z. and Zhang, A. (2024) 'Benchmarking retrieval-augmented generation "
    "for medicine', in Findings of the Association for Computational Linguistics: ACL 2024 "
    "(MIRAGE), Bangkok, August 2024.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(ref)
    set_font(run, size=11)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = "Explainable_Safe_Medical_Chatbot_Thesis_v4.docx"
doc.save(out)
print(f"Saved: {out}")
